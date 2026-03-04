from __future__ import annotations

from pathlib import Path
import sys

ROOT = Path(__file__).resolve().parents[1]
if str(ROOT) not in sys.path:
    sys.path.insert(0, str(ROOT))

from docx import Document

from cotizador_backend import generar_cotizacion_backend


def collect_docx_text(path: Path) -> str:
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        parts.append(p.text or "")
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    parts.append(p.text or "")
    for section in doc.sections:
        for p in section.header.paragraphs:
            parts.append(p.text or "")
        for p in section.footer.paragraphs:
            parts.append(p.text or "")
    return "\n".join(parts)


def main() -> None:
    out_docx = ROOT / "salidas" / "test_replacements_e2e.docx"
    result = generar_cotizacion_backend(
        modelo="CM860",
        cliente="TEST_PLACEHOLDER_CLIENT",
        ruta_salida_word=str(out_docx),
        ruta_salida_pdf="",
        precio_base_override=1000,
        precio_total_override=1200,
        opciones_overrides={
            "voltaje": "TEST_VOLT_123",
            "operacion": "TEST_OP_456",
            "bomba": "TEST_PUMP_789",
        },
        payload_replacements={
            "voltaje": "TEST_VOLT_123",
            "operacion": "TEST_OP_456",
            "bomba": "TEST_PUMP_789",
        },
    )

    docx_path = Path(result["ruta_word"])
    if not docx_path.exists():
        raise AssertionError(f"No se encontró DOCX de salida: {docx_path}")

    text = collect_docx_text(docx_path)
    expected_values = ["TEST_VOLT_123", "TEST_OP_456", "TEST_PUMP_789"]
    for value in expected_values:
        if value not in text:
            raise AssertionError(f"No se encontró valor esperado en DOCX: {value}")

    if "{{" in text or "}}" in text:
        raise AssertionError("Todavía existen placeholders sin reemplazar en el DOCX generado")

    print("OK", docx_path)


if __name__ == "__main__":
    main()
