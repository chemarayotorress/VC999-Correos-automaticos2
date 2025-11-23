"""
Backend reutilizable para generar cotizaciones Packaging+ sin depender de la GUI.
Encapsula la lógica de precios, plantillas y generación de Word/PDF en una
función central que puede ser reutilizada por la interfaz gráfica o por la CLI.
"""
from __future__ import annotations

import os
import shutil
import unicodedata
from datetime import datetime
from decimal import Decimal
from typing import Any, Dict, List, Optional, Tuple

try:
    from docx import Document
except Exception:  # pragma: no cover - dependencia opcional
    Document = None  # type: ignore

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:  # pragma: no cover - dependencia opcional
    docx2pdf_convert = None  # type: ignore

from machine_catalog import load_catalog
from template_mapping import TemplateMappingManager

APP_TITLE = "VC999 Packaging+ Cotizador"
DEFAULT_CONCEPTS = [("Con orden de compra", "35"), ("Contra aviso de entrega", "55"), ("Al instalar", "10")]

OPTION_TRANSLATIONS = {
    "Automatic lid, WITH mechanical cut": "Tapa automática CON corte mecánico",
    "Automatic lid with NO mechanical cut": "Tapa automática SIN corte mecánico",
    "Bi-active Sealing System": "Sistema de sellado biactivo",
    "Gas Flush": "Descarga de gas",
    "Positive Air Sealer": "Sellador de aire positivo",
    "Lid size": "Altura de tapa",
    "Pump Options": "Opciones de bomba",
    "Operation": "Operación",
    "Voltage": "Voltaje",
    "Machine Direction": "Dirección de la máquina",
    "Product Width (mm)": "Ancho del producto (mm)",
    "Product Height (mm)": "Altura del producto (mm)",
    "Product Length (mm)": "Longitud del producto (mm)",
    "Reject System": "Sistema de rechazo",
    "NOM-001-SCFI-2018/2014 Certification": "Certificación NOM-001-SCFI-2018/2014",
    "Sample parts kit included": "Kit de piezas de muestra incluido",
    "Index": "Índice",
    "Tray unload system": "Sistema de descarga de bandeja",
    "Tray unload": "Descarga de bandeja",
    "Registro de fotografías": "Registro de fotografías",
    "Proceso": "Proceso",
    "Configuración de matriz": "Configuración de matriz",
    "Configuración de la matriz": "Configuración de la matriz",
    "La forma de la matriz (Geometría)": "La forma de la matriz (Geometría)",
    "Forma de la matriz": "Forma de la matriz",
    "Yes": "Sí",
    "No": "No",
    "None": "Ninguno",
}

_OPTION_TO_SP = sorted(OPTION_TRANSLATIONS.items(), key=lambda kv: len(kv[0]), reverse=True)
_OPTION_TO_EN = sorted(((v, k) for k, v in OPTION_TRANSLATIONS.items()), key=lambda kv: len(kv[0]), reverse=True)


def _app_dir() -> str:
    try:
        return os.path.dirname(os.path.abspath(__file__))
    except Exception:  # pragma: no cover - defensa
        return os.getcwd()


def _sanitize_filename(text: str) -> str:
    if not isinstance(text, str):
        return "archivo"
    cleaned = [
        ch if ch.isalnum() or ch in ("-", "_") else "_"
        for ch in text.strip()
    ]
    safe = "".join(cleaned).strip("_")
    if not safe:
        safe = "archivo"
    return safe[:80]


def _fmt_money(val: Decimal, currency: str = "USD") -> str:
    try:
        s = f"{Decimal(val):,.2f}"
    except Exception:
        s = "0.00"
    return f"US${s}" if currency.upper() == "USD" else f"$ {s} {currency.upper()}"


def _parse_decimal_safe(text: Any) -> Decimal:
    if text is None:
        return Decimal("0")
    t = str(text).strip()
    if not t:
        return Decimal("0")
    t = t.replace(",", "")
    try:
        return Decimal(t)
    except Exception:
        return Decimal("0")


def _to_decimal(value: Any) -> Decimal:
    try:
        if isinstance(value, Decimal):
            return value
        return Decimal(str(value).replace(',', ''))
    except Exception:
        return Decimal('0')


def _normalize_option_data(data: Any):
    if isinstance(data, dict) and data.get('type') == 'checkbox':
        price = _to_decimal(data.get('price', 0))
        return ('chk', price)
    if isinstance(data, dict) and data.get('type') == 'select':
        choices = []
        for choice in data.get('choices', []):
            label = str(choice.get('label', ''))
            price = _to_decimal(choice.get('price', 0))
            choices.append((label, price))
        return choices
    if isinstance(data, tuple) and len(data) == 2 and data[0] == 'chk':
        return ('chk', _to_decimal(data[1]))
    if isinstance(data, (list, tuple)):
        choices = []
        for item in data:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                label = str(item[0])
                price = _to_decimal(item[1])
                choices.append((label, price))
        return choices
    return []


def _normalize_machine_catalog(raw_catalog: Dict[str, Dict]) -> Dict[str, Dict[str, Any]]:
    catalog: Dict[str, Dict[str, Any]] = {}
    for template, data in raw_catalog.items():
        base = _to_decimal((data or {}).get('base', 0))
        normalized = {'base': base, 'options': {}}
        for name, opt_data in (data or {}).get('options', {}).items():
            normalized['options'][name] = _normalize_option_data(opt_data)
        catalog[template] = normalized
    return catalog


def _apply_option_translation(text: str, lang: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    result = text
    if lang == "es":
        for en, es in _OPTION_TO_SP:
            result = result.replace(en, es)
    else:
        for es, en in _OPTION_TO_EN:
            result = result.replace(es, en)
    return result


def _to_spanish_ui(text: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    return _apply_option_translation(text, "es")


def _normalize_key(text: str) -> str:
    nk = unicodedata.normalize("NFKD", str(text)).encode("ascii", "ignore").decode("ascii")
    nk = nk.lower().replace(" ", "").replace("_", "").replace("-", "")
    return nk


def _to_bool_or_none(value: Any) -> Optional[bool]:
    if value is None:
        return None
    if isinstance(value, bool):
        return value
    text = str(value).strip().lower()
    if text in {"1", "true", "si", "sí", "yes", "on"}:
        return True
    if text in {"0", "false", "no", "off"}:
        return False
    return bool(text)


def _replace_in_paragraph(paragraph, key: str, value: str):
    if key not in paragraph.text:
        return
    runs = paragraph.runs
    full = "".join(r.text for r in runs)
    new = full.replace(key, value)
    if not runs:
        paragraph.add_run(new)
        return
    runs[0].text = new
    for r in runs[1:]:
        r.text = ""


def _replace_in_table(table, key: str, value: str):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, key, value)


def docx_replace_placeholders(doc: "Document", mapping: dict):
    for p in doc.paragraphs:
        for k, v in mapping.items():
            _replace_in_paragraph(p, k, v)
    for table in doc.tables:
        for k, v in mapping.items():
            _replace_in_table(table, k, v)


def _convert_docx_to_pdf(input_path: str, output_path: str) -> str:
    try:
        out_dir = os.path.dirname(output_path) or "."
        os.makedirs(out_dir, exist_ok=True)
        import subprocess

        subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", out_dir, input_path],
            check=True,
            stdout=subprocess.DEVNULL,
            stderr=subprocess.DEVNULL,
        )
        gen = os.path.join(out_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
        if os.path.abspath(gen) != os.path.abspath(output_path):
            try:
                os.replace(gen, output_path)
            except Exception:
                shutil.copyfile(gen, output_path)
                os.remove(gen)
        return output_path if os.path.exists(output_path) else ""
    except Exception:
        return ""


def _build_option_summary(conf_options: Dict[str, Any], overrides: Dict[str, Any]) -> Tuple[Dict[str, str], Decimal]:
    selected: Dict[str, str] = {}
    total = Decimal("0")
    for name, data in conf_options.items():
        override_key = _normalize_key(name)
        override_val = None
        for k, v in overrides.items():
            if _normalize_key(k) == override_key:
                override_val = v
                break
        if isinstance(data, tuple) and data[0] == "chk":
            price = data[1]
            val_bool = bool(override_val) if override_val is not None else False
            if val_bool:
                total += price
            selected[name.lower()] = "Sí" if val_bool else "No"
        else:
            # select type
            choices: List[Tuple[str, Decimal]] = list(data) if isinstance(data, (list, tuple)) else []
            choice_value = choices[0][0] if choices else ""
            if override_val is not None:
                for label, _price in choices:
                    if _normalize_key(label) == _normalize_key(str(override_val)) or _normalize_key(label.split("($")[0]) == _normalize_key(str(override_val)):
                        choice_value = label
                        break
            # add price of selected label
            for label, price in choices:
                if label == choice_value:
                    total += price
                    break
            clean = choice_value.split("($")[0].strip()
            selected[name.lower()] = _to_spanish_ui(clean)
    return selected, total


def _apply_template_mapping(kind: str, template_name: str, data: Dict[str, Any], context: Dict[str, Any]):
    mapping = TemplateMappingManager.load_mapping(kind, template_name)
    for placeholder, conf in mapping.items():
        if not isinstance(conf, dict):
            continue
        mode = conf.get("mode")
        if mode == "field":
            value = context.get(conf.get("value"))
        elif mode == "text":
            value = conf.get("value", "")
        else:
            value = None
        if value is not None:
            data[placeholder] = str(value)


# ------------------ API principal ------------------
def generar_cotizacion_backend(
    modelo: str,
    cliente: str,
    asesor: Optional[str] = None,
    fecha: Optional[str] = None,
    validez_dias: Optional[int] = None,
    moneda: Optional[str] = None,
    flete_texto: Optional[str] = None,
    flete_monto: Optional[float] = None,
    notas: Optional[str] = None,
    opciones_overrides: Optional[dict] = None,
    ruta_salida_word: Optional[str] = None,
    ruta_salida_pdf: Optional[str] = None,
) -> dict:
    """Genera la cotización (DOCX y opcional PDF) usando la lógica de Packaging+.

    No depende de la GUI y devuelve un diccionario con los datos relevantes.
    """

    if not modelo:
        raise ValueError("Se requiere un modelo")
    template_name = modelo if modelo.lower().endswith(".docx") else f"{modelo}.docx"

    catalog = _normalize_machine_catalog(load_catalog())
    conf = catalog.get(template_name)
    if conf is None:
        raise ValueError(f"Modelo no encontrado en catálogo: {template_name}")

    base = conf.get("base", Decimal("0"))
    overrides = opciones_overrides or {}
    selected, opt_total = _build_option_summary(conf.get("options", {}), overrides)
    total = base + opt_total

    # Datos generales
    today_str = datetime.now().strftime("%d/%m/%Y")
    fecha_valor = fecha or today_str
    asesor_valor = asesor or ""
    moneda_valor = (moneda or "USD").upper()
    validez_valor = f"{validez_dias} días" if validez_dias is not None else "30 días"
    notas_valor = notas or ""
    flete_texto_val = flete_texto or ""
    flete_monto_val = _parse_decimal_safe(flete_monto) if flete_monto is not None else Decimal("0")

    # Contrato comercial
    contrato = []
    for idx, default in enumerate(DEFAULT_CONCEPTS, start=1):
        pct_key = f"contrato{idx}_porcentaje"
        cond_key = f"contrato{idx}_condicion"
        pct_val = overrides.get(pct_key)
        cond_val = overrides.get(cond_key)
        porcentaje = str(pct_val if pct_val is not None else default[1]).replace("%", "")
        condicion = cond_val if cond_val is not None else default[0]
        contrato.append({"porcentaje": porcentaje, "condicion": condicion})

    context = {
        "cliente": cliente,
        "fecha": fecha_valor,
        "asesor": asesor_valor,
        "validez": validez_valor,
        "disponibilidad": "En stock",
        "precio_base": _fmt_money(base, moneda_valor),
        "precio_total": _fmt_money(total, moneda_valor),
        "base_numeric": float(base),
        "total_numeric": float(total),
        "options": selected,
        "conceptos": contrato,
        "conceptos_resumen": ", ".join(
            f"{c['porcentaje']}% - {c['condicion']}" for c in contrato if c.get("porcentaje")
        ),
    }

    def _put(mapping: dict, names, value):
        for n in names:
            mapping[f"{{{{{n}}}}}"] = str(value)

    data: Dict[str, str] = {}
    _put(data, ["cliente", "nombre del cliente", "nombre_del_cliente"], cliente)
    _put(data, ["fecha"], fecha_valor)
    _put(data, ["asesor"], asesor_valor)
    _put(data, ["disponibilidad"], context["disponibilidad"])
    _put(data, ["validez"], validez_valor)
    _put(data, ["precio", "total"], context["precio_total"])
    _put(data, ["precio_base"], context["precio_base"])
    _put(data, ["notas"], notas_valor)
    _put(data, ["moneda"], moneda_valor)
    _put(data, ["flete_texto"], flete_texto_val)
    if flete_monto_val:
        _put(data, ["flete_monto"], _fmt_money(flete_monto_val, moneda_valor))

    def sel(*keys):
        for k in keys:
            v = selected.get(k.lower())
            if v:
                return v
        return ""

    _put(data, ["voltage", "voltaje"], sel("Voltage", "Voltaje"))
    _put(data, ["lid size", "lid_size", "altura_tapa", "tamano_tapa", "tamaño_tapa"], sel("Lid size"))
    _put(data, ["pump options", "pump_options", "opcion_bomba", "bomba", "bomba_de_vacio"], sel("Pump Options"))
    vop = sel("Operation", "Operacion") or "No"
    vop_lower = vop.lower()
    if any(keyword in vop_lower for keyword in ("with mechanical cut", "con corte mecánico", "con corte mecanico", "sí", "si", "yes")):
        corte_val = "Con corte mecánico"
    elif any(keyword in vop_lower for keyword in ("no mechanical cut", "sin corte mecánico", "sin corte mecanico", "no", "none", "ninguno")):
        corte_val = "Sin corte mecánico"
    else:
        corte_val = "Sin corte mecánico"
    _put(data, ["operation", "operacion"], vop)
    _put(data, ["corte_mecanico"], corte_val)

    for k in list(selected.keys()):
        if "gas flush" in k or "inyeccion de gas" in k or "inyección de gas" in k:
            _put(data, ["gas_flush", "gas flush", "descarga_gas", "inyeccion_gas", "inyección_gas"], selected[k])
            break
    for k in list(selected.keys()):
        if "positive air sealer" in k:
            _put(data, ["positive_air", "positive air", "positive_air_sealer"], selected[k])
            break
    for k in list(selected.keys()):
        if "bi-active sealing system" in k or "bi active sealing system" in k:
            _put(data, ["sellado_biactivo", "bi-active sealing system", "bi_active_sealing_system"], selected[k])
            break

    for i, c in enumerate(contrato, start=1):
        pct = str(c.get("porcentaje", "")).strip()
        condicion = _to_spanish_ui(str(c.get("condicion", "")).strip())
        if pct and not pct.endswith("%"):
            pct = f"{pct}%"
        _put(data, [f"concepto{i}", f"concept{i}", f"concept_{i}"], pct)
        _put(data, [f"vencimiento{i}", f"fecha_vencimiento{i}", f"vence{i}", f"due{i}"], condicion)

    # Mapeos adicionales TS
    _put(data, ["photo_registration", "registro_fotografias", "registro_de_fotografias"], sel("Registro de fotografías", "Registro de fotograf\u00edas"))
    _put(data, ["die_configuration", "configuracion_matriz", "configuración_matriz"], sel("Configuración de matriz", "Configuración de la matriz"))
    _put(data, ["die_shape", "forma_matriz"], sel("La forma de la matriz (Geometría)", "Forma de la matriz"))
    _put(data, ["tipo_empaque"], sel("Proceso", "La forma de la matriz (Proceso)"))

    idx_val = None
    for k, v in selected.items():
        if "índice" in k or "indice" in k or "index" in k:
            idx_val = v
            break
    _put(data, ["index", "índice", "indice"], idx_val or "")

    tray_val = None
    for k, v in selected.items():
        if "descarga de bandeja" in k or "tray unload" in k or "bandeja fácil" in k or "bandeja facil" in k:
            tray_val = v
            break
    _put(data, ["sistema_descarga_bandeja", "tray_unload_system", "tray unload system", "sistema_de_descarga_de_bandeja"], tray_val or "")

    _apply_template_mapping('packaging', template_name, data, context)

    if Document is None:
        raise RuntimeError("python-docx no está instalado")

    tpl_path = os.path.join(_app_dir(), template_name)
    if not os.path.exists(tpl_path):
        raise FileNotFoundError(f"No se encontró la plantilla {template_name}")

    doc = Document(tpl_path)
    docx_replace_placeholders(doc, data)

    os.makedirs(os.path.dirname(ruta_salida_word or os.path.join(_app_dir(), "salidas")), exist_ok=True)
    default_name = f"Cotizacion_{_sanitize_filename(modelo)}_{_sanitize_filename(cliente)}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
    out_docx = ruta_salida_word or os.path.join(_app_dir(), "salidas", default_name)
    out_pdf = ruta_salida_pdf
    doc.save(out_docx)

    if ruta_salida_pdf is not None:
        out_pdf = ruta_salida_pdf
    elif ruta_salida_word:
        out_pdf = ruta_salida_word.replace(".docx", ".pdf")

    pdf_path = ""
    if out_pdf:
        try:
            if docx2pdf_convert:
                docx2pdf_convert(out_docx, out_pdf)
                pdf_path = out_pdf
            else:
                pdf_path = _convert_docx_to_pdf(out_docx, out_pdf)
        except Exception:
            pdf_path = ""

    try:
        backup_dir = os.path.join(_app_dir(), "respaldos")
        os.makedirs(backup_dir, exist_ok=True)
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        base_doc = os.path.splitext(os.path.basename(out_docx))[0]
        b_docx = os.path.join(backup_dir, f"{base_doc}__{ts}.docx")
        shutil.copy2(out_docx, b_docx)
        if pdf_path and os.path.exists(pdf_path):
            base_pdf = os.path.splitext(os.path.basename(pdf_path))[0]
            b_pdf = os.path.join(backup_dir, f"{base_pdf}__{ts}.pdf")
            shutil.copy2(pdf_path, b_pdf)
    except Exception:
        pass

    return {
        "modelo": modelo,
        "cliente": cliente,
        "asesor": asesor_valor,
        "fecha": fecha_valor,
        "validez_dias": validez_dias if validez_dias is not None else 30,
        "moneda": moneda_valor,
        "total": float(total),
        "ruta_word": out_docx,
        "ruta_pdf": pdf_path if pdf_path else "",
        "flete_texto": flete_texto_val,
        "flete_monto": float(flete_monto_val or 0),
        "opciones": selected,
        "contrato": contrato,
    }


def generar_cotizacion_desde_json(datos: dict) -> str:
    """Genera una cotización a partir de un diccionario y devuelve la ruta del PDF."""

    if not isinstance(datos, dict):
        raise ValueError("Se esperaba un diccionario de datos")

    modelo = (datos.get("modelo") or "").strip()
    cliente = (datos.get("nombre_cliente") or datos.get("cliente") or "").strip()
    if not modelo:
        raise ValueError("El campo 'modelo' es obligatorio")
    if not cliente:
        raise ValueError("El campo 'nombre_cliente' es obligatorio")

    asesor = datos.get("asesor")
    validez_dias = datos.get("validez_dias")
    moneda = datos.get("tipo_moneda") or datos.get("moneda")
    notas = datos.get("notas")
    flete_texto = datos.get("flete_texto")
    flete_monto = datos.get("flete_monto")

    overrides: Dict[str, Any] = {}

    def _set_override(key: str, value: Any) -> None:
        if value is None:
            return
        if isinstance(value, str) and value.strip() == "":
            return
        overrides[key] = value

    _set_override("altura_tapa", datos.get("altura_tapa") or datos.get("numero_tapa"))
    _set_override("operacion", datos.get("operacion"))
    _set_override("opcion_bomba", datos.get("opcion_bomba"))
    _set_override("kit_muestras", datos.get("kit_muestras"))
    _set_override("kit_muestras_nit", datos.get("kit_muestras_nit"))

    for key, raw in (
        ("descarga_gas", datos.get("inyeccion_gas")),
        ("sistema_biactivo", datos.get("sistema_biactivo")),
        ("aire_positivo", datos.get("aire_positivo")),
    ):
        bval = _to_bool_or_none(raw)
        if bval is not None:
            overrides[key] = bval

    for idx in range(1, 4):
        _set_override(f"contrato{idx}_porcentaje", datos.get(f"contrato{idx}_porcentaje"))
        _set_override(f"contrato{idx}_condicion", datos.get(f"contrato{idx}_condicion"))

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = f"Cotizacion_{_sanitize_filename(modelo)}_{_sanitize_filename(cliente)}_{timestamp}"
    salida_dir = os.path.join(_app_dir(), "salidas")
    os.makedirs(salida_dir, exist_ok=True)
    ruta_word = os.path.join(salida_dir, f"{base_name}.docx")
    ruta_pdf = os.path.join(salida_dir, f"{base_name}.pdf")

    resultado = generar_cotizacion_backend(
        modelo=modelo,
        cliente=cliente,
        asesor=asesor,
        validez_dias=int(validez_dias) if validez_dias is not None else None,
        moneda=moneda,
        flete_texto=flete_texto,
        flete_monto=flete_monto,
        notas=notas,
        opciones_overrides=overrides,
        ruta_salida_word=ruta_word,
        ruta_salida_pdf=ruta_pdf,
    )

    pdf_path = resultado.get("ruta_pdf") or ruta_pdf
    if not pdf_path or not os.path.exists(pdf_path):
        raise FileNotFoundError(f"No se pudo generar el PDF en {pdf_path}")

    return pdf_path

