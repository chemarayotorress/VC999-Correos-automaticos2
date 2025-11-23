#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
VC999 Packaging+ Cotizador
Versión: 6.3.29
Cambios vs 6.3.18:
- CM1100: tapa 12" fuerza bomba "300 m3" sin bloquear el combobox.
- Exclusión Gas Flush con corte mecánico en CM780/860/900A/1100 (consolidado).
- TS455: opción "Proceso" presente; TS540 sin "Proceso".
- Panel de opciones se autoajusta a su contenido (menos espacio en blanco).
- Selector de idioma aplicado a más textos estáticos.
- Materials: moneda predeterminada MXN.
- Oculta consola en Windows.
"""

import os, sys, json, uuid, traceback, shutil, subprocess
from typing import Dict, Any, List, Tuple
from machine_catalog import load_catalog as load_machine_catalog, save_catalog as save_machine_catalog, MachineCatalogEditor
from template_mapping import TemplateMappingManager
from cotizador_backend import generar_cotizacion_backend

# === I18N helpers ===
LANG_DEFAULT = "es"  # siempre iniciar en español
_I18N_OPT = {
    # Canonical English -> Spanish
    "Gas Flush": "Inyección de gas",
    "Positive Air Sealer": "Selladora de aire positivo",
    "Bi-active sealing system": "Sistema de sellado bi-activo",
    "Bi active sealing system": "Sistema de sellado bi-activo",
    "With mechanical cut": "Con corte mecánico",
    "with mechanical cut": "con corte mecánico",
    "Without mechanical cut": "Sin corte mecánico",
    "NO mechanical Cut": "Sin corte mecánico",
    "Mechanical Cut": "Corte mecánico",
    "Operation": "Operación",
    "Voltage": "Voltaje",
    "Lid size": "Altura de tapa",
    "Die Configuration": "Configuración de la matriz",
    "Die Shape": "Forma de la matriz",
    "Photo Registration": "Registro de fotografías",
    "Index": "Índice",
    "Tray unload system": "Sistema de descarga de bandeja fácil",
    "Tray Unload System": "Sistema de descarga de bandeja fácil",
    "Automatic lid": "Tapa automática",
    "Manual lid": "Tapa manual",
    "Skin": "Skin",
    "MAP/Lidding": "MAP/Sellado",
    "MAP / Lidding": "MAP/Sellado",
    "None": "Ninguno",
    "Yes": "Sí",
    "No": "No",
}
# Tokens to translate inside composed options
_I18N_TOKENS = sorted(_I18N_OPT.keys(), key=len, reverse=True)

def _opt_to_es(text: str) -> str:
    if not isinstance(text, str): return text
    out = text
    for k in _I18N_TOKENS:
        out = out.replace(k, _I18N_OPT[k])
    return out

# Build reverse map (Spanish -> English canonical)
_I18N_OPT_REV = {v: k for k, v in _I18N_OPT.items()}

def _opt_to_en(text: str) -> str:
    if not isinstance(text, str): return text
    out = text
    for k, v in _I18N_OPT_REV.items():
        out = out.replace(k, v)
    return out
# === End I18N helpers ===
from datetime import datetime
from decimal import Decimal, InvalidOperation

import tkinter as tk
from tkinter import ttk, messagebox, filedialog
# --- Early crash logger ---
import sys as _sys_e, traceback as _tb_e, os as _os_e
def _boot_excepthook(et, ev, tb):
    try: appdir=_os_e.path.dirname(_os_e.path.abspath(_sys_e.argv[0]))
    except: appdir=_os_e.getcwd()
    for name in ("error.log","error.txt"):
        try:
            with open(_os_e.path.join(appdir,name),"a",encoding="utf-8") as f:
                f.write("["+__import__("datetime").datetime.now().strftime("%Y-%m-%d %H:%M:%S")+"] "+repr(ev)+"\n")
                f.write("".join(_tb_e.format_exception(et,ev,tb))+"\n\n")
        except: pass
try: _sys_e.excepthook=_boot_excepthook
except: pass


# Dependencias opcionales
try:
    from docx import Document
except Exception:
    Document = None

try:
    from docx2pdf import convert as docx2pdf_convert
except Exception:
    docx2pdf_convert = None

APP_TITLE = "VC999 Packaging+ Cotizador v6.3.29"

# ----- Utilidades de ruta y configuración -----
def _app_dir() -> str:
    try:
        return os.path.dirname(os.path.abspath(sys.argv[0]))
    except Exception:
        return os.getcwd()

CFG_FILE = "app_config.json"
HIST_PACK = "historial_packaging.json"
HIST_MATS = "historial_materials.json"
HISTORY_DIR = "historial_docs"
BACKUP_DIR = "respaldos"
ADVISORS_FILE = "advisors.json"
DEFAULT_ADVISORS = ["José Manuel Rayotorres Martínez"]
MATERIALS_TEMPLATE_NAME = "Cotizacion Materials.docx"


def _is_materials_template(name: str) -> bool:
    try:
        return name.strip().lower() == MATERIALS_TEMPLATE_NAME.lower()
    except Exception:
        return False


def _sanitize_filename(text: str) -> str:
    if not isinstance(text, str):
        return "archivo"
    cleaned = [
        ch if ch.isalnum() or ch in ("-", "_") else "_"
        for ch in text.strip()
    ]
    safe = "".join(cleaned).strip("_")
    if not safe:
        safe = "archivo"
    return safe[:80]

def _read_cfg():
    p = os.path.join(_app_dir(), CFG_FILE)
    if not os.path.exists(p):
        return {}
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}

def _write_cfg(cfg: dict):
    try:
        with open(os.path.join(_app_dir(), CFG_FILE), "w", encoding="utf-8") as f:
            json.dump(cfg, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def _load_json_list(path: str, default_list):
    p = os.path.join(_app_dir(), path)
    if not os.path.exists(p):
        return list(default_list)
    try:
        with open(p, "r", encoding="utf-8") as f:
            data = json.load(f)
            if isinstance(data, list):
                return data
    except Exception:
        pass
    return list(default_list)

def _save_json_list(path: str, data_list):
    p = os.path.join(_app_dir(), path)
    try:
        with open(p, "w", encoding="utf-8") as f:
            json.dump(data_list, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def _load_hist(path: str) -> list:
    p = os.path.join(_app_dir(), path)
    if not os.path.exists(p):
        return []
    try:
        with open(p, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return []

def _save_hist(path: str, rows: list) -> None:
    p = os.path.join(_app_dir(), path)
    try:
        with open(p, "w", encoding="utf-8") as f:
            json.dump(rows, f, ensure_ascii=False, indent=2)
    except Exception:
        pass

def _write_error_log(e: Exception):
    try:
        with open(os.path.join(_app_dir(), "error.log"), "a", encoding="utf-8") as f:
            f.write(f"[{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}] {repr(e)}\n")
            f.write(traceback.format_exc() + "\n\n")
    except Exception:
        pass

# ----- Utilidades numéricas y formato -----
def _fmt_money(val: Decimal, currency: str = "USD") -> str:
    try:
        s = f"{Decimal(val):,.2f}"
    except Exception:
        s = "0.00"
    return f"US${s}" if currency.upper() == "USD" else f"$ {s} {currency.upper()}"

def _parse_decimal_safe(text: str) -> Decimal:
    if text is None:
        return Decimal("0")
    t = str(text).strip()
    if not t:
        return Decimal("0")
    t = t.replace(",", "")
    try:
        return Decimal(t)
    except InvalidOperation:
        try:
            if t.count(".") > 1:
                parts = t.split(".")
                t2 = "".join(parts[:-1]) + "." + parts[-1]
                return Decimal(t2)
        except Exception:
            pass
    return Decimal("0")

def _pct_clean(s: str) -> str:
    try:
        t = (s or "").replace("%","").strip()
        v = float(t)
        v = 0.0 if v < 0 else v
        v = 100.0 if v > 100 else v
        return f"{int(v)}%" if v.is_integer() else f"{v:.2f}%"
    except Exception:
        return "0%"

# ------------------ configuración de máquinas ------------------

def _to_decimal(value: Any) -> Decimal:
    try:
        if isinstance(value, Decimal):
            return value
        return Decimal(str(value).replace(',', ''))
    except Exception:
        return Decimal('0')


def _normalize_option_data(data: Any):
    if isinstance(data, dict) and data.get('type') == 'checkbox':
        price = _to_decimal(data.get('price', 0))
        return ('chk', price)
    if isinstance(data, dict) and data.get('type') == 'select':
        choices = []
        for choice in data.get('choices', []):
            label = str(choice.get('label', ''))
            price = _to_decimal(choice.get('price', 0))
            choices.append((label, price))
        return choices
    if isinstance(data, tuple) and len(data) == 2 and data[0] == 'chk':
        return ('chk', _to_decimal(data[1]))
    if isinstance(data, (list, tuple)):
        choices = []
        for item in data:
            if isinstance(item, (list, tuple)) and len(item) >= 2:
                label = str(item[0])
                price = _to_decimal(item[1])
                choices.append((label, price))
        return choices
    return []


def _normalize_machine_catalog(raw_catalog: Dict[str, Dict]) -> Dict[str, Dict[str, Any]]:
    catalog: Dict[str, Dict[str, Any]] = {}
    for template, data in raw_catalog.items():
        base = _to_decimal((data or {}).get('base', 0))
        normalized = {'base': base, 'options': {}}
        for name, opt_data in (data or {}).get('options', {}).items():
            normalized['options'][name] = _normalize_option_data(opt_data)
        catalog[template] = normalized
    return catalog

DEFAULT_CONCEPTS = [("Con orden de compra", "35"), ("Contra aviso de entrega", "55"), ("Al instalar", "10")]

# Traducciones de UI (clave ES -> {es,en})
TRANSLATIONS = {
    "Eliminar": {"es": "Eliminar", "en": "Delete"},
    "Cancelar": {"es": "Cancelar", "en": "Cancel"},
    "Abrir PDF": {"es": "Abrir PDF", "en": "Open PDF"},
    "Abrir DOCX": {"es": "Abrir DOCX", "en": "Open DOCX"},
    "Abrir archivo": {"es": "Abrir archivo", "en": "Open file"},
    "Notas": {"es": "Notas", "en": "Notes"},
    "Monto": {"es": "Monto", "en": "Amount"},
    "Flete": {"es": "Flete", "en": "Freight"},
    "Moneda": {"es": "Moneda", "en": "Currency"},
    "Eliminar último": {"es": "Eliminar último", "en": "Remove last"},
    "Agregar ítem": {"es": "Agregar ítem", "en": "Add item"},
    "Unitario": {"es": "Unitario", "en": "Unit price"},
    "Cantidad": {"es": "Cantidad", "en": "Quantity"},
    "Descripción": {"es": "Descripción", "en": "Description"},
    "Ítems": {"es": "Ítems", "en": "Items"},
    "Datos de la cotización": {"es": "Datos de la cotización", "en": "Quotation data"},
    "Opciones de máquina": {"es": "Opciones de máquina", "en": "Machine options"},
    "Contrato comercial": {"es": "Contrato comercial", "en": "Commercial contract"},
    "Limpiar": {"es": "Limpiar", "en": "Clear"},
    "Generar Word/PDF": {"es": "Generar Word/PDF", "en": "Generate Word/PDF"},
    "Encabezado": {"es": "Encabezado", "en": "Header"},
    "Nuevo ítem": {"es": "Nuevo ítem", "en": "New item"},
    "Opciones": {"es": "Opciones", "en": "Options"},
    "Flete y notas": {"es": "Flete y notas", "en": "Freight and notes"},
    "Subtotal:": {"es": "Subtotal:", "en": "Subtotal:"},
    "IVA:": {"es": "IVA:", "en": "VAT:"},
    "Total:": {"es": "Total:", "en": "Total:"},
    "Packaging+": {"es": "Packaging+", "en": "Packaging+"},
    "Materials": {"es": "Materials", "en": "Materials"},
    "Historial": {"es": "Historial", "en": "History"},
    "Cliente:": {"es": "Cliente:", "en": "Client:"},
    "Fecha:": {"es": "Fecha:", "en": "Date:"},
    "Asesor:": {"es": "Asesor:", "en": "Advisor:"},
    "Validez:": {"es": "Validez:", "en": "Validity:"},
    "Plantilla (.docx):": {"es": "Plantilla (.docx):", "en": "Template (.docx):"},
    "Plantilla:": {"es": "Plantilla:", "en": "Template:"},
    "Precio base (USD):": {"es": "Precio base (USD):", "en": "Base price (USD):"},
    "Disponibilidad:": {"es": "Disponibilidad:", "en": "Availability:"},
    "Concepto (%)": {"es": "Concepto (%)", "en": "Concept (%)"},
    "Fecha de vencimiento": {"es": "Fecha de vencimiento", "en": "Due date"},
    "IVA (%)": {"es": "IVA (%)", "en": "VAT (%)"},
    "Agregar": {"es": "Agregar", "en": "Add"},
    "Buscar:": {"es": "Buscar:", "en": "Search:"},
    "Filtrar": {"es": "Filtrar", "en": "Filter"},
    "Abrir...": {"es": "Abrir...", "en": "Open..."},
    "Exportar JSON": {"es": "Exportar JSON", "en": "Export JSON"},
    "Refrescar": {"es": "Refrescar", "en": "Refresh"},
    "Producto": {"es": "Producto", "en": "Product"},
    "Precio unitario": {"es": "Precio unitario", "en": "Unit price"},
    "Buscar plantilla": {"es": "Buscar plantilla", "en": "Browse template"},
    "Métricas": {"es": "Métricas", "en": "Metrics"},
    "En stock": {"es": "En stock", "en": "In stock"},
    "De 8 a 6 semanas": {"es": "De 8 a 6 semanas", "en": "6 to 8 weeks"},
    "No aplica": {"es": "No aplica", "en": "Not applicable"},
    "Incluido": {"es": "Incluido", "en": "Included"},
    "No incluido": {"es": "No incluido", "en": "Not included"},
    "Con orden de compra": {"es": "Con orden de compra", "en": "With purchase order"},
    "Contra aviso de entrega": {"es": "Contra aviso de entrega", "en": "Upon delivery notice"},
    "Al instalar": {"es": "Al instalar", "en": "Upon installation"},
}

HIST_HEADERS = {
    "fecha": {"es": "Fecha", "en": "Date"},
    "cliente": {"es": "Cliente", "en": "Client"},
    "plantilla": {"es": "Plantilla", "en": "Template"},
    "monto": {"es": "Monto", "en": "Amount"},
    "docx": {"es": "DOCX", "en": "DOCX"},
    "pdf": {"es": "PDF", "en": "PDF"},
}

MDM_DETECTOR_TEMPLATES = {"Detector de Metales MDM4121.docx"}

def _get_lang_code(value: str) -> str:
    v = (value or "").strip().lower()
    if v.startswith("en"):
        return "en"
    return "es"


def _resolve_translation(text: str):
    if not isinstance(text, str):
        return None
    for key, mapping in TRANSLATIONS.items():
        if text == mapping.get("es") or text == mapping.get("en"):
            return mapping
    return None


def _translate_text(text: str, lang: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    mapping = _resolve_translation(text)
    if mapping:
        return mapping.get(lang, text)
    return text


OPTION_TRANSLATIONS = {
    "Automatic lid, WITH mechanical cut": "Tapa automática CON corte mecánico",
    "Automatic lid with NO mechanical cut": "Tapa automática SIN corte mecánico",
    "Bi-active Sealing System": "Sistema de sellado biactivo",
    "Gas Flush": "Descarga de gas",
    "Positive Air Sealer": "Sellador de aire positivo",
    "Lid size": "Altura de tapa",
    "Pump Options": "Opciones de bomba",
    "Operation": "Operación",
    "Voltage": "Voltaje",
    "Machine Direction": "Dirección de la máquina",
    "Product Width (mm)": "Ancho del producto (mm)",
    "Product Height (mm)": "Altura del producto (mm)",
    "Product Length (mm)": "Longitud del producto (mm)",
    "Reject System": "Sistema de rechazo",
    "NOM-001-SCFI-2018/2014 Certification": "Certificación NOM-001-SCFI-2018/2014",
    "Sample parts kit included": "Kit de piezas de muestra incluido",
    "Index": "Índice",
    "Tray unload system": "Sistema de descarga de bandeja",
    "Tray unload": "Descarga de bandeja",
    "Registro de fotografías": "Registro de fotografías",
    "Proceso": "Proceso",
    "Configuración de matriz": "Configuración de matriz",
    "Configuración de la matriz": "Configuración de la matriz",
    "La forma de la matriz (Geometría)": "La forma de la matriz (Geometría)",
    "Forma de la matriz": "Forma de la matriz",
    "Yes": "Sí",
    "No": "No",
    "None": "Ninguno",
}

_OPTION_TO_SP = sorted(OPTION_TRANSLATIONS.items(), key=lambda kv: len(kv[0]), reverse=True)
_OPTION_TO_EN = sorted(((v, k) for k, v in OPTION_TRANSLATIONS.items()), key=lambda kv: len(kv[0]), reverse=True)


def _apply_option_translation(text: str, lang: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    result = text
    if lang == "es":
        for en, es in _OPTION_TO_SP:
            result = result.replace(en, es)
    else:
        for es, en in _OPTION_TO_EN:
            result = result.replace(es, en)
    return result


def _to_spanish_ui(text: str) -> str:
    if not isinstance(text, str) or not text:
        return text
    mapping = _resolve_translation(text)
    if mapping:
        return mapping.get("es", text)
    return _apply_option_translation(text, "es")

# ------------------ utilidades DOCX ------------------
def _replace_in_paragraph(paragraph, key: str, value: str):
    if key not in paragraph.text:
        return
    runs = paragraph.runs
    full = "".join(r.text for r in runs)
    new = full.replace(key, value)
    if not runs:
        paragraph.add_run(new); return
    runs[0].text = new
    for r in runs[1:]:
        r.text = ""

def _replace_in_table(table, key: str, value: str):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                _replace_in_paragraph(p, key, value)

def docx_replace_placeholders(doc: "Document", mapping: dict):
    for p in doc.paragraphs:
        for k, v in mapping.items():
            _replace_in_paragraph(p, k, v)
    for table in doc.tables:
        for k, v in mapping.items():
            _replace_in_table(table, k, v)

# ------------------ UI principal ------------------
class App(tk.Tk):
    def _mat_generate(self):
        """Genera DOCX y PDF en Materials desde plantilla elegida."""
        try:
            from docx import Document

            # Helper: reemplaza SOLO {{flete_texto}} / {{flete texto}} en cualquier parte del DOCX
            def _force_replace_flete(docx_path, value):
                import zipfile, os, tempfile, shutil, re
                tmpdir = tempfile.mkdtemp()
                try:
                    with zipfile.ZipFile(docx_path, 'r') as zin:
                        zin.extractall(tmpdir)
                    pat = re.compile(
                        r'\{(?:\s|<[^>]+>)*\{' +
                        r'(?:\s|<[^>]+>)*f(?:\s|<[^>]+>)*l(?:\s|<[^>]+>)*e(?:\s|<[^>]+>)*t(?:\s|<[^>]+>)*e' +
                        r'(?:\s|<[^>]+>|_|-)+' +
                        r'(?:\s|<[^>]+>)*t(?:\s|<[^>]+>)*e(?:\s|<[^>]+>)*x(?:\s|<[^>]+>)*t(?:\s|<[^>]+>)*o' +
                        r'(?:\s|<[^>]+>)*\}(?:\s|<[^>]+>)*\}',
                        re.I
                    )
                    for root, _dirs, files in os.walk(tmpdir):
                        for fn in files:
                            if not fn.lower().endswith(".xml"):
                                continue
                            fpath = os.path.join(root, fn)
                            try:
                                xml = open(fpath, "r", encoding="utf-8").read()
                            except Exception:
                                continue
                            xml2, n = pat.subn(value, xml)
                            if n:
                                with open(fpath, "w", encoding="utf-8") as f:
                                    f.write(xml2)
                    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                        for root, _dirs, files in os.walk(tmpdir):
                            for fn in files:
                                ap = os.path.join(root, fn)
                                arc = os.path.relpath(ap, tmpdir)
                                zout.write(ap, arc)
                finally:
                    try: shutil.rmtree(tmpdir)
                    except Exception: pass

            # Helper: reemplazo XML para placeholders dentro de textboxes/runs partidos
            def _replace_xml_placeholders(docx_path, mapping):
                import zipfile, os, tempfile, shutil, re
                tmpdir = tempfile.mkdtemp()
                try:
                    with zipfile.ZipFile(docx_path, 'r') as zin:
                        zin.extractall(tmpdir)
                    for root, _dirs, files in os.walk(tmpdir):
                        for fn in files:
                            if not fn.lower().endswith(".xml"):
                                continue
                            fpath = os.path.join(root, fn)
                            try:
                                xml = open(fpath, "r", encoding="utf-8").read()
                            except Exception:
                                continue
                            changed = False
                            # exact pass
                            for k, v in mapping.items():
                                if k and k in xml:
                                    xml = xml.replace(k, v); changed = True
                            # fuzzy pass allowing tags/spaces between braces and letters
                            def _rx(ph):
                                name = ph.strip("{}")
                                parts = []
                                for ch in name:
                                    parts.append(re.escape(ch) + r'(?:\s|<[^>]+>)*')
                                inner = r'(?:\s|<[^>]+>)*'.join(parts)
                                pat = r'\{(?:\s|<[^>]+>)*\{' + inner + r'\}(?:\s|<[^>]+>)*\}'
                                return re.compile(pat, re.I)
                            for k, v in mapping.items():
                                try:
                                    rx = _rx(k)
                                    xml, n = rx.subn(v, xml)
                                    if n: changed = True
                                except Exception:
                                    pass
                            if changed:
                                with open(fpath, "w", encoding="utf-8") as f:
                                    f.write(xml)
                    with zipfile.ZipFile(docx_path, 'w', zipfile.ZIP_DEFLATED) as zout:
                        for root, _dirs, files in os.walk(tmpdir):
                            for fn in files:
                                ap = os.path.join(root, fn)
                                arc = os.path.relpath(ap, tmpdir)
                                zout.write(ap, arc)
                finally:
                    try: shutil.rmtree(tmpdir)
                    except Exception: pass
        except Exception as e:
            try:
                messagebox.showerror(APP_TITLE, f"Falta python-docx:\n{e}")
            except Exception:
                pass
            return
        try:
            import os
            from datetime import datetime
            from decimal import Decimal

            
            def _replace_placeholders_doc(doc, rep: dict):
                # Replace in paragraphs robustly (across runs)
                for p in doc.paragraphs:
                    txt = p.text
                    changed = False
                    for k, v in rep.items():
                        if k in txt:
                            txt = txt.replace(k, v); changed = True
                    if changed:
                        for r in p.runs:
                            r.text = ""
                        if p.runs:
                            p.runs[0].text = txt
                        else:
                            p.add_run(txt)
                # Replace in tables cells as whole text
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            ctext = cell.text
                            changed = False
                            for k, v in rep.items():
                                if k in ctext:
                                    ctext = ctext.replace(k, v); changed = True
                            if changed:
                                cell.text = ctext

            def _fill_template_items_table(doc, items, moneda, subtotal, iva_pct, iva, grand):
                from docx.shared import Cm, RGBColor
                from docx.enum.text import WD_ALIGN_PARAGRAPH
                from docx.oxml import OxmlElement
                from docx.oxml.ns import qn

                def _shade(cell, hex_fill):
                    tcPr = cell._tc.get_or_add_tcPr()
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), hex_fill.replace('#',''))
                    tcPr.append(shd)

                def _set_col_widths(tbl, cms):
                    tbl.autofit = False
                    for row in tbl.rows:
                        for i,w in enumerate(cms):
                            try:
                                row.cells[i].width = Cm(w)
                            except Exception:
                                pass

                # Find the table with header containing Qty, Product, Each, Price
                target = None
                for t in doc.tables:
                    if len(t.rows) == 0:
                        continue
                    hdr = t.rows[0].cells
                    head_txt = " | ".join([c.text.strip() for c in hdr])
                    if (("Qty" in head_txt or "Cantidad" in head_txt) and ("Product" in head_txt or "Descripción" in head_txt)
                        and ("Each" in head_txt or "Unitario" in head_txt)):
                        target = t
                        break
                if target is None:
                    # Crear nueva tabla con el formato solicitado
                    target = doc.add_table(rows=1, cols=4)
                    # Estilo con bordes
                    try:
                        target.style = 'Table Grid'
                    except Exception:
                        try:
                            target.style = 'TableGrid'
                        except Exception:
                            pass
                    # Encabezados
                    hdr = target.rows[0].cells
                    hdr[0].text = "Qty"; hdr[1].text = "Product"; hdr[2].text = "Each (MXN)"; hdr[3].text = "Price (MXN)"
                    # Color de encabezado rojo y texto blanco centrado
                    for c in hdr:
                        _shade(c, "#E31C24")
                        for p in c.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.bold = True
                                r.font.color.rgb = RGBColor(255,255,255)
                    # Anchos 10% | 50% | 20% | 20% aprox en cm para páginas A4/Letter
                    _set_col_widths(target, [2.0, 10.0, 4.0, 4.0])

                # Quitar filas de datos y totales si existían
                while len(target.rows) > 1:
                    r = target.rows[-1]
                    target._tbl._element.remove(r._tr)

                # Asegurar encabezado con formato si venía de plantilla
                hdr = target.rows[0].cells
                try:
                    from docx.shared import RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH
                    for c in hdr:
                        for p in c.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for r in p.runs:
                                r.font.bold = True
                                try:
                                    r.font.color.rgb = RGBColor(255,255,255)
                                except Exception:
                                    pass
                except Exception:
                    pass
                _set_col_widths(target, [2.0, 10.0, 4.0, 4.0])

                # Agregar items
                for desc, qty, each, total in items:
                    r = target.add_row().cells
                    r[0].text = str(qty)
                    r[1].text = desc
                    r[2].text = _fmt_money(each, moneda)
                    r[3].text = _fmt_money(total, moneda)

                # Totales
                labcol = 2  # tercera columna
                # Subtotal
                r = target.add_row().cells
                r[0].text = ""
                r[1].text = ""
                r[2].text = "Product Subtotal"
                r[3].text = _fmt_money(subtotal, moneda)
                # IVA
                r = target.add_row().cells
                r[0].text = ""
                r[1].text = ""
                r[2].text = f"IVA {iva_pct}%"
                r[3].text = _fmt_money(iva, moneda)
                # Total
                r = target.add_row().cells
                r[0].text = ""
                r[1].text = ""
                r[2].text = "Total"
                r[3].text = _fmt_money(grand, moneda)
                # Sombreado suave en totales y alinear a la derecha
                for row in target.rows[-3:]:
                    try:
                        _shade(row.cells[2], "#FFF2CC")
                        _shade(row.cells[3], "#FFF2CC")
                        from docx.enum.text import WD_ALIGN_PARAGRAPH
                        row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                        row.cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    except Exception:
                        pass
                return True
# 1) Plantilla
            tpl = None
            if hasattr(self, "var_m_tpl"):
                tpl = (self.var_m_tpl.get() or "").strip()
            if not tpl:
                tpl = os.path.join(_app_dir(), "Cotizacion Materials.docx")
            if not os.path.exists(tpl):
                messagebox.showerror(APP_TITLE, "No se encontró la plantilla seleccionada.\nVerifica 'Cotizacion Materials.docx' o el selector de plantilla.")
                return

            # 2) Datos encabezado
            cliente = (self.var_m_cliente.get() or "").strip() or "Cliente"
            fecha = (self.var_m_fecha.get() or "").strip()
            asesor = (self.var_m_asesor.get() or "").strip() or "José Manuel Rayotorres Martínez"
            validez = (self.var_m_validez.get() or "").strip()
            notas = (self.txt_m_notas.get("1.0","end").strip() if hasattr(self,"txt_m_notas") else "")  # usuario decide
            if not notas:
                notas = ""  # dejar espacio en blanco

            
            modo = _to_spanish_ui(self.var_m_flete_modo.get() if hasattr(self,"var_m_flete_modo") else "No aplica")
            try:
                monto = _parse_decimal_safe(self.var_m_flete_monto.get()) if hasattr(self,"var_m_flete_monto") else Decimal("0")
            except Exception:
                monto = Decimal("0")
            moneda = self.var_moneda.get() if hasattr(self,"var_moneda") else "MXN"
            if modo == "Incluido":
                flete_texto = "Flete: Incluido" if monto <= 0 else f"Flete: Incluido {_fmt_money(monto, moneda)}"
            elif modo == "No incluido":
                flete_texto = "Flete: No incluido"
            elif modo == "No aplica":
                flete_texto = ""
            else:
                flete_texto = str(modo)
# 3) Ítems
            items = []
            for row, v_desc, v_qty, v_each in getattr(self,"item_rows", []):
                desc = (v_desc.get() or "").strip()
                qty = int(_parse_decimal_safe(v_qty.get() or "0"))
                each = _parse_decimal_safe(v_each.get() or "0")
                if not desc or qty <= 0:
                    continue
                items.append((desc, qty, each, qty*each))

            # 4) Documento
            doc = Document(tpl)
            rep = {
                "{{nombre del cliente}}": cliente,
                "{{cliente}}": cliente,
                "{{fecha}}": fecha,
                "{{asesor}}": asesor,
                "{{validez}}": validez,
                "{{notas}}": notas,
                "{{flete_texto}}": flete_texto,
            }
            if items:
                _i_desc, _i_qty, _i_each, _i_total = items[0]
                rep.update({
                    "{{qty1}}": str(_i_qty),
                    "{{product1}}": _i_desc,
                    "{{each1}}": _fmt_money(_i_each, moneda),
                    "{{price1}}": _fmt_money(_i_total, moneda),
                })
            for p in doc.paragraphs:
                for k,v in rep.items():
                    if k in p.text:
                        for r in p.runs:
                            r.text = r.text.replace(k, v)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for k,v in rep.items():
                            if k in cell.text:
                                cell.text = cell.text.replace(k, v)

            
            # 5) Tabla + totales (usar la tabla existente de la plantilla)
            subtotal = Decimal("0")
            iva_pct = int(_parse_decimal_safe(self.var_iva.get())) if hasattr(self,"var_iva") else 16
            # Calcular totales
            for _d, _q, _e, _t in items:
                subtotal += Decimal(str(_t))
            iva = (subtotal * Decimal(iva_pct) / Decimal("100")).quantize(Decimal("0.01"))
            grand = (subtotal + iva).quantize(Decimal("0.01"))
            context = self._gather_materials_context(os.path.basename(tpl), cliente, fecha, asesor, validez, notas, flete_texto, moneda, items, subtotal, Decimal(iva_pct), iva, grand)

            # Insertar filas en la tabla que contiene {{qty1}} y {{product1}}
            target = None; row_idx = None
            for t in doc.tables:
                for i, row in enumerate(t.rows):
                    row_txt = " | ".join(c.text for c in row.cells)
                    if "{{qty1}}" in row_txt and "{{product1}}" in row_txt:
                        target = t; row_idx = i; break
                if target is not None: break

            
            # Si no se encontró por placeholders, buscar por la fila que ya contiene el primer ítem
            if target is None and items:
                _d0, _q0, _e0, _t0 = items[0]
                _q0s = str(_q0)
                for t in doc.tables:
                    for i, row in enumerate(t.rows):
                        row_txt = " | ".join(c.text for c in row.cells)
                        if _q0s in row_txt and _d0 in row_txt:
                            target = t; row_idx = i; break
                    if target is not None: break
            if target is not None and items:
                # Primera fila
                _d, _q, _e, _t = items[0]
                cells = target.rows[row_idx].cells
                cells[0].text = str(_q)
                cells[1].text = _d
                cells[2].text = _fmt_money(_e, moneda)
                cells[3].text = _fmt_money(_t, moneda)

                # clonar base e insertar filas extra debajo
                from copy import deepcopy
                from docx.table import _Row
                prev_tr = target.rows[row_idx]._tr
                for _d, _q, _e, _t in items[1:]:
                    clone_tr = deepcopy(prev_tr)
                    prev_tr.addnext(clone_tr)
                    new_row = _Row(clone_tr, target)
                    new_row.cells[0].text = str(_q)
                    new_row.cells[1].text = _d
                    new_row.cells[2].text = _fmt_money(_e, moneda)
                    new_row.cells[3].text = _fmt_money(_t, moneda)
                    prev_tr = clone_tr

            # Actualizar placeholders de totales en todo el documento
            rep.update({
                "{{subtotal}}": _fmt_money(subtotal, moneda),
                "{{iva}}": _fmt_money(iva, moneda),
                "{{total}}": _fmt_money(grand, moneda),
            })
            self._apply_template_mapping("materials", os.path.basename(tpl), rep, context)
            for p in doc.paragraphs:
                for k,v in rep.items():
                    if k in p.text:
                        for r in p.runs:
                            r.text = r.text.replace(k, v)
            for t in doc.tables:
                for row in t.rows:
                    for cell in row.cells:
                        for k,v in rep.items():
                            if k in cell.text:
                                cell.text = cell.text.replace(k, v)
# 6) Guardar DOCX
            out_name = f"{cliente} {fecha}.docx".replace("/", "-")
            init_dir = os.path.dirname(tpl) if tpl else _app_dir()
            out_docx = filedialog.asksaveasfilename(defaultextension=".docx",
                                                    filetypes=[("Word","*.docx")],
                                                    initialdir=init_dir, initialfile=out_name,
                                                    title="Guardar cotización")
            if not out_docx:
                messagebox.showinfo(APP_TITLE, "Operación cancelada.")
                return
            os.makedirs(os.path.dirname(out_docx) or ".", exist_ok=True)
            doc.save(out_docx)
            _force_replace_flete(out_docx, flete_texto)
            extra_xml = {"{{notas}}": notas, "{{flete_texto}}": flete_texto, "{{subtotal}}": _fmt_money(subtotal, moneda), "{{iva}}": _fmt_money(iva, moneda), "{{total}}": _fmt_money(grand, moneda)}
            for key, value in rep.items():
                if isinstance(key, str) and key.startswith("{{"):
                    extra_xml[key] = value
            _replace_xml_placeholders(out_docx, extra_xml)


            # 7) PDF
            out_pdf = out_docx.replace(".docx",".pdf")
            try:
                if docx2pdf_convert:
                    try:
                        docx2pdf_convert(out_docx, out_pdf)
                    except Exception:
                        _ = self._convert_docx_to_pdf(out_docx, out_pdf)
                else:
                    _ = self._convert_docx_to_pdf(out_docx, out_pdf)
            except Exception as e:
                _write_error_log(e)

            
            # 7.2) Respaldos automáticos de DOCX y PDF
            try:
                import os, shutil
                from datetime import datetime
                backup_dir = os.path.join(_app_dir(), "respaldos")
                os.makedirs(backup_dir, exist_ok=True)
                ts = datetime.now().strftime("%Y%m%d_%H%M%S")
                # DOCX
                try:
                    base_doc = os.path.splitext(os.path.basename(out_docx))[0]
                    b_docx = os.path.join(backup_dir, f"{base_doc}__{ts}.docx")
                    shutil.copy2(out_docx, b_docx)
                except Exception as _e:
                    _write_error_log(_e)
                # PDF si existe
                try:
                    if os.path.exists(out_pdf):
                        base_pdf = os.path.splitext(os.path.basename(out_pdf))[0]
                        b_pdf = os.path.join(backup_dir, f"{base_pdf}__{ts}.pdf")
                        shutil.copy2(out_pdf, b_pdf)
                except Exception as _e:
                    _write_error_log(_e)
            except Exception as _e:
                _write_error_log(_e)
# 8) Historial
            try:
                stored_docx, stored_pdf = self._history_store_files(
                    "materials",
                    out_docx,
                    out_pdf if os.path.exists(out_pdf) else "",
                )
                history_entry = {
                    "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    "cliente": cliente,
                    "plantilla": os.path.basename(tpl),
                    "monto": f"{_fmt_money(grand, moneda)}",
                    "docx": stored_docx or out_docx,
                    "pdf": stored_pdf or (out_pdf if os.path.exists(out_pdf) else ""),
                    "total_numeric": float(grand),
                    "original_docx": out_docx,
                    "original_pdf": out_pdf if os.path.exists(out_pdf) else "",
                }
                self._history_add("materials", history_entry)
            except Exception as e:
                _write_error_log(e)

            # 9) Mensaje final
            msg_pdf = out_pdf if os.path.exists(out_pdf) else "(PDF no disponible)"
            messagebox.showinfo(APP_TITLE, "Documento generado:\n" + out_docx + "\n" + msg_pdf)

        except Exception as e:
            _write_error_log(e)
            messagebox.showerror(APP_TITLE, f"Ocurrió un error al generar:\n{e}")

    # --- Historial: UI y utilidades ---
    def _hist_build_common(self, parent, is_pack: bool):
        import tkinter as tk
        from tkinter import ttk, messagebox
        self._hist_is_pack = getattr(self, "_hist_is_pack", {})
        self._hist_is_pack[parent] = is_pack
        top = ttk.Frame(parent); top.pack(fill="both", expand=True, padx=8, pady=8)
        # Barra de búsqueda
        self._hist_search = getattr(self, "_hist_search", {})
        self._hist_is_pack = getattr(self, "_hist_is_pack", {})
        self._hist_is_pack[parent] = is_pack
        import tkinter as tk
        search_bar = ttk.Frame(top); search_bar.pack(fill="x", pady=(0,6))
        lbl_search = ttk.Label(search_bar, text="Buscar:")
        lbl_search.pack(side="left")
        self._register_translated_widget("Buscar:", lbl_search)
        qvar = tk.StringVar(value="")
        self._hist_search[parent] = qvar
        ent = ttk.Entry(search_bar, textvariable=qvar, width=40)
        ent.pack(side="left", padx=6)
        btn_filter = ttk.Button(search_bar, text="Filtrar", command=lambda tv=None,par=parent: self._hist_refresh(tv_ref, par))
        btn_filter.pack(side="left")
        self._register_translated_widget("Filtrar", btn_filter)
        btn_clear = ttk.Button(search_bar, text="Limpiar", command=lambda: (qvar.set(""), self._hist_refresh(tv_ref, parent)))
        btn_clear.pack(side="left", padx=6)
        self._register_translated_widget("Limpiar", btn_clear)
        cols = ("fecha","cliente","plantilla","monto","docx","pdf")
        tv = ttk.Treeview(top, columns=cols, show="headings", height=12)
        tv_ref = tv
        self._hist_tree_kind = getattr(self, "_hist_tree_kind", {})
        self._hist_tree_kind[tv] = "packaging" if is_pack else "materials"
        lang_code = _get_lang_code(self.lang_var.get() if hasattr(self, "lang_var") else "Español")
        for c in cols:
            mapping = HIST_HEADERS.get(c)
            text = mapping.get(lang_code, mapping.get("es", c.upper())) if mapping else c.upper()
            tv.heading(c, text=text)
        tv.column("fecha", width=140, anchor="w")
        tv.column("cliente", width=180, anchor="w")
        tv.column("plantilla", width=160, anchor="w")
        tv.column("monto", width=110, anchor="e")
        tv.column("docx", width=260, anchor="w")
        tv.column("pdf", width=260, anchor="w")
        tv.pack(fill="both", expand=True, side="top")
        self._hist_trees.append((tv, cols))
        # Doble clic o Enter abre PDF si existe, si no DOCX
        tv.bind("<Double-1>", lambda e, tv=tv: self._hist_open_default(tv))
        tv.bind("<Return>", lambda e, tv=tv: self._hist_open_default(tv))
        btns = ttk.Frame(top); btns.pack(fill="x", pady=6)
        btn_refresh = ttk.Button(btns, text="Refrescar", command=lambda tv=tv,par=parent: self._hist_refresh(tv, par))
        btn_refresh.pack(side="left")
        self._register_translated_widget("Refrescar", btn_refresh)
        btn_pdf = ttk.Button(btns, text="Abrir PDF", command=lambda tv=tv: self._hist_open_selected(tv, which="pdf"))
        btn_pdf.pack(side="left", padx=6)
        self._register_translated_widget("Abrir PDF", btn_pdf)
        btn_del = ttk.Button(btns, text="Eliminar", command=lambda tv=tv,par=parent: self._hist_delete_selected(tv, par))
        btn_del.pack(side="right")
        self._register_translated_widget("Eliminar", btn_del)
        self._hist_refresh(tv, parent)

    def _hist_refresh(self, tree, parent):
        trees = [tree] if tree else []
        if not trees:
            try:
                for child in parent.winfo_children():
                    for sub in child.winfo_children():
                        if isinstance(sub, ttk.Treeview):
                            trees.append(sub)
            except Exception:
                pass
        kind = "packaging" if self._hist_is_pack.get(parent, True) else "materials"
        rows = self._history_fetch(kind)
        q = ""
        try:
            qvar = getattr(self, "_hist_search", {}).get(parent)
            q = (qvar.get() if qvar else "").strip().lower()
        except Exception:
            q = ""
        if q:
            filtered = []
            for row in rows:
                if any(q in str(row.get(k, "")).lower() for k in ("fecha", "cliente", "plantilla", "monto", "docx", "pdf")):
                    filtered.append(row)
            rows = filtered
        for tv in trees:
            for item in tv.get_children():
                tv.delete(item)
            for row in rows:
                values = (
                    row.get("fecha", ""),
                    row.get("cliente", ""),
                    row.get("plantilla", ""),
                    row.get("monto", ""),
                    row.get("docx", ""),
                    row.get("pdf", ""),
                )
                tv.insert("", "end", iid=row.get("id"), values=values)
        self._update_metrics_panel()


    def _hist_open_default(self, tree):
        if not tree:
            return
        sel = tree.selection()
        if not sel:
            return
        item_id = sel[0]
        kind = self._hist_tree_kind.get(tree, "packaging")
        row = None
        for candidate in self._hist_cache.get(kind, []):
            if candidate.get("id") == item_id:
                row = candidate
                break
        values = tree.item(item_id, "values") if tree.exists(item_id) else ()
        pdf_path = (row or {}).get("pdf", "") or (row or {}).get("original_pdf", "")
        docx_path = (row or {}).get("docx", "") or (row or {}).get("original_docx", "")
        if not pdf_path and values and len(values) > 5:
            pdf_path = values[5]
        if not docx_path and values and len(values) > 4:
            docx_path = values[4]
        target = "pdf" if pdf_path else "docx"
        self._hist_open_selected(tree, which=target)


    def _hist_open_selected(self, tree, which="docx"):
        import os, glob
        from tkinter import messagebox
        sel = tree.selection()
        if not sel:
            return
        item_id = sel[0]
        kind = self._hist_tree_kind.get(tree, "packaging")
        rows = self._hist_cache.get(kind, [])
        row = next((r for r in rows if r.get("id") == item_id), None)
        vals = tree.item(item_id, "values") if tree.exists(item_id) else ()
        if not row and vals:
            row = {
                "docx": vals[4] if len(vals) > 4 else "",
                "pdf": vals[5] if len(vals) > 5 else "",
            }
        path = (row or {}).get(which, "").strip()
        if path and os.path.exists(path):
            self._open_path(path)
            return
        fallback_key = "original_pdf" if which.lower() == "pdf" else "original_docx"
        if row and fallback_key in row:
            alt = str(row.get(fallback_key, "")).strip()
            if alt and os.path.exists(alt):
                self._open_path(alt)
                return
        base_source = path or ((row or {}).get("docx") or ((row or {}).get("pdf") or ""))
        base = os.path.splitext(os.path.basename(base_source))[0]
        base = base.split("__")[0] if "__" in base else base
        backups = self._history_collect_backups(base)
        cand = ""
        if backups:
            preferred_ext = ("." + which.lower()) if which else ""
            if preferred_ext:
                for fp in backups:
                    if fp.lower().endswith(preferred_ext):
                        cand = fp
                        break
            cand = cand or backups[0]
        if cand:
            self._open_path(cand)
        else:
            messagebox.showwarning(APP_TITLE, f"No hay ruta de {which.upper()} y no se halló respaldo.")

    
    def _hist_delete_selected(self, tree, parent):
        sel = tree.selection()
        if not sel:
            return
        item_id = sel[0]
        kind = self._hist_tree_kind.get(tree, "packaging")
        rows = self._hist_cache.get(kind, [])
        row = next((r for r in rows if r.get("id") == item_id), None)
        vals = tree.item(item_id, "values") if tree.exists(item_id) else ()
        if not row and vals:
            row = {
                "docx": vals[4] if len(vals) > 4 else "",
                "pdf": vals[5] if len(vals) > 5 else "",
            }
        base_pdf = os.path.splitext(os.path.basename((row or {}).get("pdf") or ""))[0]
        base_doc = os.path.splitext(os.path.basename((row or {}).get("docx") or ""))[0]
        base = base_pdf or base_doc
        base = base.split("__")[0] if "__" in base else base
        try:
            self._history_delete(kind, item_id)
            try:
                import os
                for key in ("docx", "pdf", "original_docx", "original_pdf"):
                    target = (row or {}).get(key)
                    if target and self._is_history_copy(target) and os.path.exists(target):
                        try:
                            os.remove(target)
                        except Exception:
                            pass
                for fp in self._history_collect_backups(base):
                    try:
                        os.remove(fp)
                    except Exception:
                        pass
            except Exception:
                pass
            self._hist_refresh(tree, parent)
        except Exception:
            from tkinter import messagebox
            messagebox.showerror(APP_TITLE, "No se pudo eliminar del historial.")

    def _open_path(self, path: str):
        import os, subprocess, sys
        try:
            if os.name == "nt":
                os.startfile(path)  # type: ignore[attr-defined]
            elif sys.platform == "darwin":
                subprocess.run(["open", path], check=False)
            else:
                subprocess.run(["xdg-open", path], check=False)
        except Exception:
            pass

    def _build_hist(self, parent):
        stats = ttk.LabelFrame(parent, text="Métricas")
        stats.pack(fill="x", padx=8, pady=(8, 4))
        self._register_translated_widget("Métricas", stats)
        kinds = [("packaging", "Packaging+"), ("materials", "Materials")]
        for kind, _label in kinds:
            self._history_kind_dir(kind)
        for idx, (kind, label) in enumerate(kinds):
            ttk.Label(stats, text=label + ":").grid(row=idx, column=0, sticky="w", padx=6, pady=2)
            var = tk.StringVar(value="-")
            self._metrics_vars[kind] = var
            ttk.Label(stats, textvariable=var).grid(row=idx, column=1, sticky="w", padx=6, pady=2)
        stats.columnconfigure(1, weight=1)

        nb = ttk.Notebook(parent); nb.pack(fill="both", expand=True)
        tab_p = ttk.Frame(nb); tab_m = ttk.Frame(nb)
        nb.add(tab_p, text="Packaging+"); nb.add(tab_m, text="Materials")
        self._hist_build_common(tab_p, is_pack=True)
        self._hist_build_common(tab_m, is_pack=False)
        self.after(500, self._update_metrics_panel)

    def __init__(self, cfg: Dict[str, Any]):
        self.cfg = cfg or {}
        self.machine_catalog_raw = load_machine_catalog()
        self.machine_catalog = _normalize_machine_catalog(self.machine_catalog_raw)
        self._filter_packaging_templates()
        self._logo_img = None
        self._trans_widgets = {}
        self._trans_comboboxes = []
        self._hist_trees = []
        self._hist_cache: Dict[str, List[Dict[str, Any]]] = {"packaging": [], "materials": []}
        self._metrics_vars: Dict[str, tk.StringVar] = {}
        try:
            self._i18n_init()
        except Exception:
            pass
        super().__init__()
        self.title(APP_TITLE)
        self._min_width = 760
        self._min_height = 520
        self.minsize(self._min_width, self._min_height)

        if os.name == "nt":
            try:
                import ctypes
                hwnd = ctypes.windll.kernel32.GetConsoleWindow()
                if hwnd:
                    ctypes.windll.user32.ShowWindow(hwnd, 0)
            except Exception:
                pass

        nb = ttk.Notebook(self); nb.pack(fill="both", expand=True)
        self._notebook = nb
        self.tab_pack = ttk.Frame(nb)
        self.tab_mat = ttk.Frame(nb)
        self.tab_hist = ttk.Frame(nb)
        nb.add(self.tab_pack, text="Packaging+")
        nb.add(self.tab_mat, text="Materials")
        nb.add(self.tab_hist, text="Historial")
        nb.bind("<<NotebookTabChanged>>", lambda e: self._resize_to_tab())

        self._build_packaging(self.tab_pack)
        self._build_materials(self.tab_mat)
        self._build_hist(self.tab_hist)

        self.after(100, self._center_once)
        self.after(150, self._on_language_change)
        self.after(250, lambda: self._refresh_machine_catalog(True))
        self.after(400, self._resize_to_tab)

    def _center_once(self):
        try:
            self._resize_to_tab(center=True)
        except Exception:
            pass

    def _resize_to_tab(self, center: bool = False):
        try:
            self.update_idletasks()
            nb = getattr(self, "_notebook", None)
            current = nb.nametowidget(nb.select()) if nb else None
            extra_w = self.winfo_reqwidth() - (nb.winfo_reqwidth() if nb else 0)
            extra_h = self.winfo_reqheight() - (nb.winfo_reqheight() if nb else 0)
            extra_w = max(extra_w, 0)
            extra_h = max(extra_h, 0)
            req_w = max(self._min_width, (current.winfo_reqwidth() if current else self.winfo_reqwidth()) + extra_w)
            req_h = max(self._min_height, (current.winfo_reqheight() if current else self.winfo_reqheight()) + extra_h)
            max_w = max(self._min_width, int(self.winfo_screenwidth() * 0.9))
            max_h = max(self._min_height, int(self.winfo_screenheight() * 0.9))
            new_w = min(req_w, max_w)
            new_h = min(req_h, max_h)
            if center:
                sw, sh = self.winfo_screenwidth(), self.winfo_screenheight()
                x, y = max((sw - new_w) // 2, 0), max((sh - new_h) // 2, 0)
                self.geometry(f"{new_w}x{new_h}+{x}+{y}")
            else:
                self.geometry(f"{new_w}x{new_h}")
        except Exception:
            pass

    def _register_translated_widget(self, key: str, widget):
        if not key or widget is None:
            return
        lst = self._trans_widgets.setdefault(key, [])
        if widget not in lst:
            lst.append(widget)

    def _register_translated_combobox(self, widget, var: tk.StringVar, base_values):
        try:
            values = list(base_values)
        except Exception:
            values = []
        self._trans_comboboxes.append((widget, var, values))

    def _apply_combobox_language(self):
        lang_code = _get_lang_code(self.lang_var.get() if hasattr(self, "lang_var") else "Español")
        for widget, var, base_values in list(self._trans_comboboxes):
            try:
                values = [_translate_text(v, lang_code) for v in base_values]
                widget.configure(values=values)
                current_base = _to_spanish_ui(var.get())
                if current_base not in base_values and base_values:
                    current_base = base_values[0]
                display = _translate_text(current_base, lang_code)
                if var.get() != display:
                    var.set(display)
            except Exception:
                continue

    def _update_hist_headings(self):
        lang_code = _get_lang_code(self.lang_var.get() if hasattr(self, "lang_var") else "Español")
        for tree, cols in list(self._hist_trees):
            try:
                for col in cols:
                    mapping = HIST_HEADERS.get(col)
                    if not mapping:
                        continue
                    tree.heading(col, text=mapping.get(lang_code, mapping.get("es", col)))
            except Exception:
                continue

    def _update_material_totals_labels(self):
        try:
            moneda = self.var_moneda.get() if hasattr(self, "var_moneda") else "MXN"
        except Exception:
            moneda = "MXN"
        totals = getattr(self, "_mat_totals", {})
        subtotal = totals.get("subtotal")
        iva = totals.get("iva")
        total = totals.get("total")
        lang_code = _get_lang_code(self.lang_var.get() if hasattr(self, "lang_var") else "Español")
        try:
            if subtotal is None:
                subtotal = Decimal("0")
            if iva is None:
                iva = Decimal("0")
            if total is None:
                total = subtotal + iva
        except Exception:
            subtotal = subtotal or Decimal("0")
            iva = iva or Decimal("0")
            total = total or (subtotal + iva)
        try:
            self.lbl_sub.config(text=f"{_translate_text('Subtotal:', lang_code)} {_fmt_money(subtotal, moneda)}")
            self.lbl_iva.config(text=f"{_translate_text('IVA:', lang_code)} {_fmt_money(iva, moneda)}")
            self.lbl_tot.config(text=f"{_translate_text('Total:', lang_code)} {_fmt_money(total, moneda)}")
        except Exception:
            pass

    def _filter_packaging_templates(self):
        try:
            for container in (self.machine_catalog_raw, self.machine_catalog):
                if not isinstance(container, dict):
                    continue
                for name in list(container.keys()):
                    if _is_materials_template(name):
                        container.pop(name, None)
        except Exception:
            pass

    def _packaging_template_names(self) -> List[str]:
        return [name for name in sorted(self.machine_catalog.keys()) if not _is_materials_template(name)]

    # ===== Packaging+ =====
    def _build_packaging(self, parent):
        # Encabezado
        header = ttk.Frame(parent); header.pack(fill="x", padx=8, pady=(8,0))
        self.logo_label = ttk.Label(header, text="")
        self.logo_label.pack(side="left", padx=(0,12))
        self._auto_logo()

        self.lang_var = tk.StringVar(value=self.cfg.get("lang", "Español"))
        lang_cb = ttk.Combobox(header, textvariable=self.lang_var, values=["Español", "English"], state="readonly", width=12)
        lang_cb.pack(side="right")
        lang_cb.bind("<<ComboboxSelected>>", lambda e: self._on_language_change())

        top = ttk.LabelFrame(parent, text="Datos de la cotización"); top.pack(fill="x", padx=8, pady=6)
        self._register_translated_widget("Datos de la cotización", top)

        lbl_plant = ttk.Label(top, text="Plantilla (.docx):"); lbl_plant.grid(row=0, column=0, sticky="w", padx=4, pady=3)
        self._register_translated_widget("Plantilla (.docx):", lbl_plant)
        self.var_plant = tk.StringVar()
        plant_vals = self._packaging_template_names()
        self.cb_plant = ttk.Combobox(top, textvariable=self.var_plant, values=plant_vals, state="readonly", width=26)
        self.cb_plant.grid(row=0, column=1, sticky="w", padx=4, pady=3)
        self.cb_plant.bind("<<ComboboxSelected>>", lambda e: self._on_template_change())

        lbl_fecha = ttk.Label(top, text="Fecha:"); lbl_fecha.grid(row=0, column=2, sticky="e", padx=4, pady=3)
        self._register_translated_widget("Fecha:", lbl_fecha)
        self.var_fecha = tk.StringVar(value=datetime.now().strftime("%d/%m/%Y"))
        ttk.Entry(top, textvariable=self.var_fecha, width=12).grid(row=0, column=3, sticky="w", padx=4, pady=3)

        lbl_cli = ttk.Label(top, text="Cliente:"); lbl_cli.grid(row=1, column=0, sticky="w", padx=4, pady=3)
        self._register_translated_widget("Cliente:", lbl_cli)
        self.var_cliente = tk.StringVar(); ttk.Entry(top, textvariable=self.var_cliente, width=34).grid(row=1, column=1, sticky="w", padx=4, pady=3)

        lbl_ase = ttk.Label(top, text="Asesor:"); lbl_ase.grid(row=1, column=2, sticky="e", padx=4, pady=3)
        self._register_translated_widget("Asesor:", lbl_ase)
        self.advisors = _load_json_list(ADVISORS_FILE, DEFAULT_ADVISORS)
        self.var_asesor = tk.StringVar(value=(self.advisors[0] if self.advisors else ""))
        self.cb_asesor = ttk.Combobox(top, textvariable=self.var_asesor, values=self.advisors, width=26)
        self.cb_asesor.grid(row=1, column=3, sticky="w", padx=4, pady=3)
        ttk.Button(top, text="Agregar", command=self._add_advisor).grid(row=1, column=4, sticky="w", padx=6)
        self._register_translated_widget("Agregar", self.cb_asesor.master.grid_slaves(row=1, column=4)[0])

        lbl_base = ttk.Label(top, text="Precio base (USD):"); lbl_base.grid(row=2, column=0, sticky="w", padx=4, pady=3)
        self._register_translated_widget("Precio base (USD):", lbl_base)
        self.var_base = tk.StringVar(value="0")
        ebase = ttk.Entry(top, textvariable=self.var_base, width=12)
        ebase.grid(row=2, column=1, sticky="w", padx=4, pady=3)
        ebase.bind("<FocusOut>", lambda e: (self._normalize_money(self.var_base), self._recalc_pack()))

        lbl_disp = ttk.Label(top, text="Disponibilidad:"); lbl_disp.grid(row=2, column=2, sticky="e", padx=4, pady=3)
        self._register_translated_widget("Disponibilidad:", lbl_disp)
        self.var_disp = tk.StringVar(value="En stock")
        disp_values = ["En stock", "De 8 a 6 semanas"]
        self.cmb_disp = ttk.Combobox(top, textvariable=self.var_disp, values=disp_values, width=18, state="readonly")
        self.cmb_disp.grid(row=2, column=3, sticky="w", padx=4, pady=3)
        self._register_translated_combobox(self.cmb_disp, self.var_disp, disp_values)

        lbl_val = ttk.Label(top, text="Validez:"); lbl_val.grid(row=2, column=4, sticky="e", padx=4, pady=3)
        self._register_translated_widget("Validez:", lbl_val)
        self.var_validez = tk.StringVar(value="30 días")
        ttk.Entry(top, textvariable=self.var_validez, width=12).grid(row=2, column=5, sticky="w", padx=4, pady=3)

        # Cuerpo a dos columnas
        mid = ttk.Frame(parent); mid.pack(fill="both", expand=True, padx=8, pady=4)
        mid.columnconfigure(0, weight=2)
        mid.columnconfigure(1, weight=1)

        lf = ttk.LabelFrame(mid, text="Opciones de máquina")
        self._register_translated_widget("Opciones de máquina", lf)
        lf.grid(row=0, column=0, sticky="nsew", padx=(0,6), pady=2)
        self.op_frame = ttk.Frame(lf)
        self.op_frame.pack(fill="both", expand=False)

        rf = ttk.LabelFrame(mid, text="Contrato comercial")
        self._register_translated_widget("Contrato comercial", rf)
        rf.grid(row=0, column=1, sticky="nsew", padx=(6,0), pady=2)
        lbl_conc = ttk.Label(rf, text="Concepto (%)"); lbl_conc.grid(row=0, column=0, padx=6, pady=4, sticky="w")
        self._register_translated_widget("Concepto (%)", lbl_conc)
        lbl_due = ttk.Label(rf, text="Fecha de vencimiento"); lbl_due.grid(row=0, column=1, padx=6, pady=4, sticky="w")
        self._register_translated_widget("Fecha de vencimiento", lbl_due)
        self.var_conc_pct = [tk.StringVar(value="35"), tk.StringVar(value="55"), tk.StringVar(value="10")]
        presets = [x[0] for x in DEFAULT_CONCEPTS]
        self.var_conc_venc = [tk.StringVar(value=presets[0]), tk.StringVar(value=presets[1]), tk.StringVar(value=presets[2])]
        for i in range(3):
            ttk.Entry(rf, textvariable=self.var_conc_pct[i], width=8).grid(row=i+1, column=0, padx=6, pady=3, sticky="w")
            cb_venc = ttk.Combobox(rf, textvariable=self.var_conc_venc[i], values=presets, width=28, state="readonly")
            cb_venc.grid(row=i+1, column=1, padx=6, pady=3, sticky="we")
            self._register_translated_combobox(cb_venc, self.var_conc_venc[i], presets)
        rf.columnconfigure(1, weight=1)

        bottom = ttk.Frame(parent); bottom.pack(fill="x", padx=8, pady=(6,8))
        self.var_total = tk.StringVar(value="US$0.00")
        ttk.Label(bottom, textvariable=self.var_total, font=("Segoe UI", 11, "bold")).pack(side="left")
        btn_clear = ttk.Button(bottom, text="Limpiar", command=self._pack_clear); btn_clear.pack(side="right", padx=(6,0))
        btn_gen = ttk.Button(bottom, text="Generar Word/PDF", command=self._pack_generate); btn_gen.pack(side="right")
        self._register_translated_widget("Limpiar", btn_clear)
        self._register_translated_widget("Generar Word/PDF", btn_gen)

        if plant_vals:
            self.var_plant.set(plant_vals[0]); self._on_template_change()

    def _refresh_machine_catalog(self, auto: bool = False):
        try:
            raw = load_machine_catalog()
            catalog = _normalize_machine_catalog(raw)
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"No se pudo cargar el catálogo: {exc}")
            return

        docx_files = set()
        try:
            for name in os.listdir(_app_dir()):
                if not name.lower().endswith(".docx"):
                    continue
                if name.startswith("~$"):
                    continue
                if _is_materials_template(name):
                    continue
                docx_files.add(name)
        except Exception:
            docx_files = set()

        new_templates: List[str] = []
        for name in sorted(docx_files):
            if name not in raw:
                raw[name] = {"base": "0", "options": {}}
                new_templates.append(name)
            if name not in catalog:
                catalog[name] = {"base": Decimal("0"), "options": {}}

        self.machine_catalog_raw = raw
        self.machine_catalog = catalog
        self._filter_packaging_templates()
        plant_vals = self._packaging_template_names()
        try:
            self.cb_plant.configure(values=plant_vals)
        except Exception:
            pass
        current = self.var_plant.get()
        if plant_vals and current not in plant_vals:
            self.var_plant.set(plant_vals[0])
            self._on_template_change()

        if new_templates:
            if auto:
                try:
                    save_machine_catalog(raw)
                except Exception:
                    pass
            elif messagebox and messagebox.askyesno(
                APP_TITLE,
                "Se detectaron nuevas plantillas. ¿Deseas completar sus datos antes de guardar?",
            ):
                def _save(data):
                    save_machine_catalog(data)
                    self.machine_catalog_raw = data
                    self.machine_catalog = _normalize_machine_catalog(data)
                    self._refresh_machine_catalog()

                MachineCatalogEditor(self, self.machine_catalog_raw, _save)
            else:
                try:
                    save_machine_catalog(raw)
                except Exception as exc:
                    if messagebox:
                        messagebox.showwarning(
                            APP_TITLE,
                            f"No se pudo guardar el catálogo actualizado:\
{exc}",
                        )

    def _open_machine_editor(self):
        def _save(data):
            save_machine_catalog(data)
            self.machine_catalog_raw = data
            self.machine_catalog = _normalize_machine_catalog(data)
            self._refresh_machine_catalog()
        MachineCatalogEditor(self, self.machine_catalog_raw, _save)

    def _ensure_history_entry(self, kind: str, entry: Dict[str, Any]) -> Dict[str, Any]:
        entry = dict(entry)
        entry.setdefault('id', uuid.uuid4().hex)
        entry.setdefault('kind', kind)
        return entry

    def _history_local_path(self, kind: str) -> str:
        return HIST_PACK if kind == 'packaging' else HIST_MATS

    def _history_kind_dir(self, kind: str) -> str:
        base_dir = os.path.join(_app_dir(), HISTORY_DIR)
        subdir = 'packaging' if kind == 'packaging' else 'materials'
        path = os.path.join(base_dir, subdir)
        try:
            os.makedirs(path, exist_ok=True)
        except Exception:
            pass
        return path

    def _is_history_copy(self, path: str) -> bool:
        if not path:
            return False
        try:
            base = os.path.abspath(os.path.join(_app_dir(), HISTORY_DIR))
            return os.path.commonpath([os.path.abspath(path), base]) == base
        except Exception:
            return False

    def _history_store_files(self, kind: str, docx_path: str, pdf_path: str) -> Tuple[str, str]:
        stored_docx = ""
        stored_pdf = ""
        try:
            base_dir = self._history_kind_dir(kind)
            ts = datetime.now().strftime("%Y%m%d_%H%M%S")
            source = docx_path or pdf_path or f"{kind}_cotizacion"
            base_name = _sanitize_filename(os.path.splitext(os.path.basename(source))[0])
            if docx_path and os.path.exists(docx_path):
                dest_docx = os.path.join(base_dir, f"{base_name}__{ts}.docx")
                shutil.copy2(docx_path, dest_docx)
                stored_docx = dest_docx
            if pdf_path and os.path.exists(pdf_path):
                dest_pdf = os.path.join(base_dir, f"{base_name}__{ts}.pdf")
                shutil.copy2(pdf_path, dest_pdf)
                stored_pdf = dest_pdf
        except Exception as exc:
            _write_error_log(exc)
        return stored_docx, stored_pdf

    def _history_local_add(self, kind: str, entry: Dict[str, Any]):
        entry = self._ensure_history_entry(kind, entry)
        rows = _load_hist(self._history_local_path(kind))
        rows.append(entry)
        _save_hist(self._history_local_path(kind), rows)

    def _history_local_delete(self, kind: str, entry_id: str):
        rows = _load_hist(self._history_local_path(kind))
        rows = [r for r in rows if r.get('id') != entry_id]
        _save_hist(self._history_local_path(kind), rows)

    def _history_add(self, kind: str, entry: Dict[str, Any]):
        entry = self._ensure_history_entry(kind, entry)
        self._history_local_add(kind, entry)
        self._hist_cache.setdefault(kind, []).append(entry)

    def _history_row_sort_key(self, row: Dict[str, Any]) -> float:
        """Return a comparable value for history rows.

        Prefers the parsed ``fecha`` field when available. If it cannot be
        parsed, fall back to ``total_numeric``. For backup rows without these
        fields, derive the value from the filesystem modification time.
        """
        if not isinstance(row, dict):
            return 0.0

        fecha = row.get("fecha")
        if isinstance(fecha, str) and fecha.strip():
            for fmt in ("%Y-%m-%d %H:%M:%S", "%Y-%m-%d %H:%M", "%Y-%m-%d"):
                try:
                    return datetime.strptime(fecha.strip(), fmt).timestamp()
                except Exception:
                    continue

        total_numeric = row.get("total_numeric")
        if total_numeric is not None:
            try:
                return float(total_numeric)
            except Exception:
                pass

        for key in ("path", "docx", "pdf"):
            path = row.get(key)
            if isinstance(path, str) and path:
                try:
                    return float(os.path.getmtime(path))
                except Exception:
                    continue

        return 0.0

    def _history_fetch(self, kind: str) -> List[Dict[str, Any]]:
        rows = _load_hist(self._history_local_path(kind))
        ensured = []
        for row in rows:
            ensured.append(self._ensure_history_entry(kind, dict(row)))
        backups = self._history_collect_backups(kind)
        backup_rows: List[Dict[str, Any]] = []
        for path in backups:
            if not isinstance(path, str):
                continue
            try:
                mtime = os.path.getmtime(path)
            except Exception:
                continue
            entry: Dict[str, Any] = {
                "id": uuid.uuid5(uuid.NAMESPACE_URL, os.path.abspath(path)).hex,
                "kind": kind,
                "fecha": datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M"),
                "total_numeric": mtime,
                "cliente": "Respaldo automático",
                "plantilla": os.path.basename(path),
                "monto": "",
                "docx": path if path.lower().endswith(".docx") else "",
                "pdf": path if path.lower().endswith(".pdf") else "",
                "path": path,
                "is_backup": True,
            }
            if not entry["docx"] and not entry["pdf"]:
                entry["docx"] = path
            backup_rows.append(self._ensure_history_entry(kind, entry))
        combined = ensured + backup_rows
        combined.sort(key=self._history_row_sort_key, reverse=True)
        self._hist_cache[kind] = combined
        return combined

    def _history_delete(self, kind: str, entry_id: str):
        if not entry_id:
            return
        self._history_local_delete(kind, entry_id)
        cache = self._hist_cache.get(kind, [])
        self._hist_cache[kind] = [r for r in cache if r.get('id') != entry_id]

    def _history_collect_backups(self, base: str) -> List[str]:
        if not base:
            return []
        try:
            import glob
            bdir = os.path.join(_app_dir(), BACKUP_DIR)
            if not os.path.isdir(bdir):
                return []
            patterns = [
                os.path.join(bdir, f"{base}__*.pdf"),
                os.path.join(bdir, f"{base}__*.docx"),
            ]
            found: List[str] = []
            for pat in patterns:
                found.extend(glob.glob(pat))
            found.sort(key=lambda p: os.path.getmtime(p), reverse=True)
            return found
        except Exception:
            return []

    def _update_metrics_panel(self):
        if not self._metrics_vars:
            return
        for kind, var in self._metrics_vars.items():
            rows = list(self._hist_cache.get(kind, []) or [])
            if not rows:
                var.set("Sin registros locales")
                continue
            total = Decimal("0")
            for row in rows:
                val = row.get("total_numeric")
                if val is None:
                    val = row.get("total")
                try:
                    total += Decimal(str(val))
                except Exception:
                    continue
            currency = "USD"
            for row in rows:
                monto = str(row.get("monto", "")).upper()
                if "MXN" in monto:
                    currency = "MXN"
                    break
                if "USD" in monto or "US$" in monto:
                    currency = "USD"
                    break
            formatted_total = _fmt_money(total, currency)
            count = len(rows)
            label = "cotización" if count == 1 else "cotizaciones"
            var.set(f"{count} {label} locales / {formatted_total}")

    def _auto_logo(self):
        path = None
        try:
            p = self.cfg.get("logo_path")
            if p and os.path.exists(p):
                path = p
            else:
                cand = os.path.join(_app_dir(), "vc999_logo.png")
                if os.path.exists(cand):
                    path = cand
                    self.cfg["logo_path"] = cand
                    _write_cfg(self.cfg)
        except Exception:
            path = None
        if not path:
            return
        try:
            img = tk.PhotoImage(file=path)
            try: img = img.subsample(2,2)
            except Exception: pass
            self._logo_img = img
            self.logo_label.configure(image=img)
        except Exception as e:
            messagebox.showwarning(APP_TITLE, f"No se pudo mostrar el logo ({e}).")

    def _add_advisor(self):
        name = (self.var_asesor.get() or "").strip()
        if not name:
            messagebox.showwarning(APP_TITLE, "Escribe el nombre del asesor y presiona 'Agregar asesor'."); return
        if name not in self.advisors:
            self.advisors.append(name)
            _save_json_list(ADVISORS_FILE, self.advisors)
            self.cb_asesor["values"] = self.advisors
            messagebox.showinfo(APP_TITLE, f"Asesor agregado: {name}")

    def _on_template_change(self):
        for w in self.op_frame.winfo_children():
            w.destroy()
        plant = self.var_plant.get()
        conf = self.machine_catalog.get(plant)
        if conf is None and self.machine_catalog:
            names = self._packaging_template_names()
            plant = names[0] if names else plant
            self.var_plant.set(plant)
            conf = self.machine_catalog.get(plant)
        if conf is None:
            conf = {"base": Decimal("0"), "options": {}}
        self.var_base.set(str(conf.get("base", Decimal("0"))))
        self.opt_vars = {}
        self.opt_widgets = {}
        self.opt_grid_info = {}
        self.opt_label_widgets = {}
        self.opt_label_grid_info = {}
        r = 0
        for name, data in conf.get("options", {}).items():
            if isinstance(data, tuple) and data[0] == "chk":
                var = tk.BooleanVar(value=False)
                self.opt_vars[name] = ("chk", var, data[1])
                chk = tk.Checkbutton(self.op_frame, text=self._tr_opt(name), variable=var, anchor="w", wraplength=520, justify="left", command=lambda n=name: (self._recalc_pack(), self._recompute_option_canvas_height()))
                chk.grid(row=r, column=0, columnspan=2, sticky="w", padx=6, pady=1)
                self.opt_widgets[name] = chk
                self.opt_grid_info[name] = chk.grid_info().copy()
                r += 1
            else:
                lbl = ttk.Label(self.op_frame, text=self._tr_opt(name) + ":")
                lbl.grid(row=r, column=0, sticky="w", padx=6, pady=1)
                self.opt_label_widgets[name] = lbl
                self.opt_label_grid_info[name] = lbl.grid_info().copy()
                var = tk.StringVar()
                vals = [label for (label, _price) in data]
                display_vals = [self._tr_opt(x) for x in vals]
                cb = ttk.Combobox(self.op_frame, textvariable=var, values=display_vals, state="readonly", width=58)
                cb.grid(row=r, column=1, sticky="we", padx=6, pady=1)
                var.set(display_vals[0] if display_vals else "")
                cb.bind("<<ComboboxSelected>>", lambda e: (self._recalc_pack(), self._recompute_option_canvas_height()))
                self.op_frame.grid_columnconfigure(1, weight=1)
                self.opt_vars[name] = ("combo", var, data)
                self.opt_widgets[name] = cb
                self.opt_grid_info[name] = cb.grid_info().copy()
                r += 1

        self._recalc_pack()
        try:
            self._setup_dynamic_behaviors(plant)
        finally:
            self._recompute_option_canvas_height()


    def _tr_opt(self, text: str) -> str:
        if not isinstance(text, str):
            return text
        lang_code = _get_lang_code(self.lang_var.get() if hasattr(self, 'lang_var') else "Español")
        translated = _translate_text(text, lang_code)
        return _apply_option_translation(translated, lang_code)

    def _normalize_money(self, var: tk.StringVar):
        try:
            v = _parse_decimal_safe(var.get())
            var.set(f"{v:,.2f}")
        except Exception:
            var.set("0")

    def _recalc_pack(self):
        base = _parse_decimal_safe(self.var_base.get()); total = base
        for name, info in self.opt_vars.items():
            if info[0] == "chk":
                var, price = info[1], info[2]
                if var.get(): total += price
            else:
                var, data = info[1], info[2]; sel = var.get()
                for label, price in data:
                    if label == sel or self._tr_opt(label) == sel: total += price; break
        self.var_total.set(f"US${total:,.2f}")

    def _gather_packaging_context(self, base: Decimal, total: Decimal, selected: Dict[str, str]) -> Dict[str, Any]:
        conceptos = []
        for idx in range(len(self.var_conc_pct)):
            pct = self.var_conc_pct[idx].get()
            due = self.var_conc_venc[idx].get() if idx < len(self.var_conc_venc) else ""
            conceptos.append({"porcentaje": pct, "concepto": due})
        resumen = ", ".join(f"{c['porcentaje']} - {c['concepto']}" for c in conceptos if c.get('porcentaje'))
        return {
            "cliente": self.var_cliente.get(),
            "fecha": self.var_fecha.get(),
            "asesor": self.var_asesor.get(),
            "validez": self.var_validez.get(),
            "disponibilidad": _to_spanish_ui(self.var_disp.get()),
            "precio_base": f"US${base:,.2f}",
            "precio_total": f"US${total:,.2f}",
            "base_numeric": float(base),
            "total_numeric": float(total),
            "options": selected,
            "conceptos": conceptos,
            "conceptos_resumen": resumen,
        }

    def _get_mapping_value(self, key: str, context: Dict[str, Any]):
        if not key:
            return None
        if key in context:
            return context[key]
        if key == "conceptos_json":
            return json.dumps(context.get("conceptos", []), ensure_ascii=False)
        if key == "options_resumen":
            return context.get("conceptos_resumen", "")
        if key.startswith("option:"):
            name = key.split(":", 1)[1].strip().lower()
            return context.get("options", {}).get(name, "")
        return context.get(key)

    def _apply_template_mapping(self, kind: str, template_name: str, data: Dict[str, Any], context: Dict[str, Any]):
        mapping = TemplateMappingManager.load_mapping(kind, template_name)
        for placeholder, conf in mapping.items():
            if not isinstance(conf, dict):
                continue
            mode = conf.get("mode")
            if mode == "field":
                value = self._get_mapping_value(conf.get("value"), context)
            elif mode == "text":
                value = conf.get("value", "")
            else:
                value = None
            if value is not None:
                data[placeholder] = str(value)

    def _gather_materials_context(self, tpl: str, cliente: str, fecha: str, asesor: str, validez: str, notas: str,
                                  flete_texto: str, moneda: str, items: List[Any], subtotal: Decimal,
                                  iva_pct: Decimal, iva: Decimal, grand: Decimal) -> Dict[str, Any]:
        items_payload = [
            {"descripcion": desc, "cantidad": qty, "unitario": str(each), "total": str(total)}
            for desc, qty, each, total in items
        ]
        return {
            "plantilla": tpl,
            "cliente": cliente,
            "fecha": fecha,
            "asesor": asesor,
            "validez": validez,
            "notas": notas,
            "flete_texto": flete_texto,
            "moneda": moneda,
            "subtotal": f"{subtotal:,.2f}",
            "iva": f"{iva:,.2f}",
            "total": f"{grand:,.2f}",
            "subtotal_monto": float(subtotal),
            "iva_monto": float(iva),
            "total_monto": float(grand),
            "iva_pct": str(iva_pct),
            "items": items_payload,
            "items_json": json.dumps(items_payload, ensure_ascii=False),
        }


    def _recompute_option_canvas_height(self):
        try:
            self.update_idletasks()
            # No scrollbars. Let frame size to content and rely on main window layout.
            self._resize_to_tab()
        except Exception:
            pass

    def _pack_clear(self):
        self.var_cliente.set("")
        self.var_fecha.set(datetime.now().strftime("%d/%m/%Y"))
        self.var_asesor.set(self.advisors[0] if self.advisors else "")
        self.var_disp.set("En stock")
        self.var_validez.set("30 días")
        plant = self.var_plant.get()
        conf = self.machine_catalog.get(plant, {"base": Decimal("0")})
        self.var_base.set(str(conf.get("base", Decimal("0"))))
        for name, info in getattr(self, "opt_vars", {}).items():
            if info[0] == "chk":
                info[1].set(False)
            else:
                var, data = info[1], info[2]
                labels = [self._tr_opt(label) for (label, _price) in data]
                var.set(labels[0] if labels else "")
        for i, v in enumerate(DEFAULT_CONCEPTS):
            self.var_conc_pct[i].set(v[1]); self.var_conc_venc[i].set(v[0])
        self._recalc_pack(); self._recompute_option_canvas_height()
        try:
            self._apply_combobox_language()
        except Exception:
            pass

    # --- Idioma ---
    def _on_language_change(self):
        lang = self.lang_var.get() if hasattr(self, 'lang_var') else "Español"
        lang_code = _get_lang_code(lang)
        try:
            self.cfg["lang"] = lang; _write_cfg(self.cfg)
        except Exception:
            pass
        for key, widgets in self._trans_widgets.items():
            mapping = TRANSLATIONS.get(key)
            if not mapping:
                continue
            text = mapping.get(lang_code, mapping.get('es', key))
            for widget in widgets:
                try:
                    widget.configure(text=text)
                except Exception:
                    pass
        # Reconstruir opciones para aplicar traducción en listas
        try:
            self._on_template_change()
        except Exception:
            pass
        # Títulos de pestañas
        try:
            nb = self.children.get('!notebook')
            if nb:
                texts = ["Packaging+", "Materials", "Historial"]
                for i, t in enumerate(texts):
                    m = TRANSLATIONS.get(t, None)
                    if m:
                        nb.tab(i, text=m.get(lang_code, m.get('es', t)))
        except Exception:
            pass
        try:
            self._apply_combobox_language()
        except Exception:
            pass
        try:
            self._update_hist_headings()
        except Exception:
            pass
        try:
            self._update_material_totals_labels()
        except Exception:
            pass
        try:
            self._resize_to_tab()
        except Exception:
            pass

    # ------------------------- comportamientos dinámicos -------------------------
    def _setup_dynamic_behaviors(self, plant: str):
        if not hasattr(self, "opt_vars"):
            return

        def show_widget(key: str):
            w = self.opt_widgets.get(key); info = self.opt_grid_info.get(key)
            if w and info:
                try:
                    w.grid(**{k: v for k, v in info.items() if k in ("row","column","columnspan","sticky","padx","pady")})
                except Exception:
                    pass

        def hide_widget(key: str):
            w = self.opt_widgets.get(key)
            if w:
                try: w.grid_remove()
                except Exception: pass

        def regrid_label(key: str, row: int, column: int, columnspan: int = 1):
            labels = getattr(self, "opt_label_widgets", {})
            label = labels.get(key)
            if not label:
                return
            try:
                label.grid(row=row, column=column, columnspan=columnspan, sticky="w", padx=6, pady=1)
                grid_store = getattr(self, "opt_label_grid_info", None)
                if isinstance(grid_store, dict):
                    grid_store[key] = label.grid_info().copy()
            except Exception:
                pass

        def regrid_option_widget(key: str, row: int, column: int, columnspan: int = 1, sticky: str = "we"):
            widget = self.opt_widgets.get(key)
            if not widget:
                return
            try:
                widget.grid(row=row, column=column, columnspan=columnspan, sticky=sticky, padx=6, pady=1)
                self.opt_grid_info[key] = widget.grid_info().copy()
            except Exception:
                pass

        def regrid_option_pair(key: str, row: int, label_column: int, value_column: int):
            regrid_label(key, row, label_column)
            regrid_option_widget(key, row, value_column)

        def regrid_checkbox(key: str, row: int, column: int, columnspan: int = 2):
            regrid_option_widget(key, row, column, columnspan=columnspan, sticky="w")

        def configure_mechanical_cut(operation_key: str, gas_key: str, positive_air_key: str = None):
            op_info = self.opt_vars.get(operation_key)
            gas_info = self.opt_vars.get(gas_key)
            pa_info = self.opt_vars.get(positive_air_key) if positive_air_key else None
            if not op_info or not gas_info:
                return
            # operation combo
            op_var = op_info[1]
            op_data = op_info[2] if len(op_info) > 2 else []
            gas_var = gas_info[1]
            pa_var = pa_info[1] if pa_info else None

            def is_mech_cut(sel: str) -> bool:
                s = (sel or "").strip().lower()
                # direct text checks in both languages
                if ("with mechanical cut" in s) or s.startswith("yes") or s.startswith("sí") or ("con corte mec" in s):
                    return True
                if ("no mechanical cut" in s) or s.startswith("none") or s.startswith("ninguno") or ("sin corte mec" in s):
                    return False
                # map back to canonical labels from data
                for lab,_ in op_data:
                    lab_l = lab.lower()
                    if sel == lab or sel == getattr(self, "_tr_opt", lambda x:x)(lab):
                        if "with mechanical cut" in lab_l or lab_l.startswith("yes"):
                            return True
                        if "no mechanical cut" in lab_l or lab_l.startswith("none"):
                            return False
                # fallback: if option key mentions mechanical cut and selection not None
                if "mechanical cut" in operation_key.lower():
                    return "none" not in s and "ninguno" not in s
                return False

            def on_op_change(*_):
                val = op_var.get()
                if is_mech_cut(val):
                    try: gas_var.set(False)
                    except Exception: pass
                    # hide gas widget
                    try:
                        w = self.opt_widgets.get(gas_key)
                        if w: w.grid_remove()
                    except Exception: pass
                    if pa_var:
                        try: pa_var.set(True)
                        except Exception: pass
                else:
                    # show gas again
                    try:
                        w = self.opt_widgets.get(gas_key); info = self.opt_grid_info.get(gas_key)
                        if w and info:
                            w.grid(**{k:v for k,v in info.items() if k in ("row","column","columnspan","sticky","padx","pady")})
                    except Exception: pass
                self._recalc_pack(); self._recompute_option_canvas_height()

            try: op_var.trace_add("write", on_op_change)
            except Exception: pass
            on_op_change()

        def configure_pump_on_lid(lid_key: str, pump_key: str, required_value_substr: str):
            lid_info = self.opt_vars.get(lid_key)
            pump_info = self.opt_vars.get(pump_key)
            if not lid_info or not pump_info:
                return
            lid_var = lid_info[1]
            pump_var = pump_info[1]
            pump_data = pump_info[2]

            default_val = pump_data[0][0] if pump_data else ""
            req_val = default_val
            for lbl, _ in pump_data:
                if "2 x 200 m3" in lbl.lower():
                    req_val = lbl; break

            def on_lid_change(*_):
                val = (lid_var.get() or "").lower()
                if "12" in val:
                    pump_var.set(req_val)
                else:
                    pump_var.set(default_val)
                self._recalc_pack()

            try:
                lid_var.trace_add("write", on_lid_change)
            except Exception:
                pass
            on_lid_change()

        if plant == "CM780.docx":
            configure_mechanical_cut("Mechanical Cut w/ Positive Air Sealer", "Gas Flush ($ 995 USD)")
        elif plant == "CM430.docx":
            configure_mechanical_cut("Operation", "Gas Flush ($ 995 USD)")
        elif plant == "CM860.docx":
            configure_mechanical_cut("Operation", "Gas Flush ($ 995 USD)")
            configure_mechanical_cut("Operation", "Gas Flush ($ 995 USD)")
        elif plant == "CM900A.docx":
            configure_mechanical_cut("Operation", "Gas Flush ($ 995 USD)")
            configure_mechanical_cut("Operation", "Gas Flush ($ 995 USD)")
        elif plant == "CM1100.docx":
            configure_mechanical_cut("Operation", "Gas Flush ($ 995 USD)")
            configure_pump_on_lid("Lid size", "Pump Options", "2 x 200 m3")
        elif plant in MDM_DETECTOR_TEMPLATES:
            try:
                self.op_frame.grid_columnconfigure(3, weight=1)
            except Exception:
                pass
            left_keys = ["Voltage", "Product Width (mm)", "Machine Direction"]
            right_keys = ["Product Height (mm)", "Product Length (mm)", "Reject System"]
            for idx, key in enumerate(left_keys):
                regrid_option_pair(key, idx, 0, 1)
            for idx, key in enumerate(right_keys):
                regrid_option_pair(key, idx, 2, 3)
            regrid_checkbox("NOM-001-SCFI-2018/2014 Certification", len(left_keys), 0, columnspan=2)
    # === I18N in App ===
    def _i18n_init(self):
        # estado de idioma
        try:
            self.lang = getattr(self, "lang", None) or LANG_DEFAULT
        except Exception:
            self.lang = LANG_DEFAULT
        self._prev_lang = self.lang
        # si existe una variable de UI para idioma, sincronizar
        try:
            if hasattr(self, "var_lang") and self.var_lang.get():
                self.lang = "en" if self.var_lang.get().lower().startswith("en") else "es"
        except Exception:
            pass
        self.after(500, self._i18n_watch)

    def _i18n_watch(self):
        # Observa cambios en self.var_lang sin reiniciar
        try:
            v = None
            if hasattr(self, "var_lang"):
                v = self.var_lang.get().lower().strip()
            if v in ("es", "español", "espanol"):
                new_lang = "es"
            elif v in ("en", "english"):
                new_lang = "en"
            else:
                new_lang = self.lang
            if new_lang != self.lang:
                self.lang = new_lang
                self._apply_language()
        except Exception:
            pass
        try:
            self.after(500, self._i18n_watch)
        except Exception:
            pass

    def _apply_language(self):
        # Actualiza valores visibles de combos ya seleccionados
        try:
            for name, info in getattr(self, "opt_vars", {}).items():
                if info[0] == "combo":
                    var, data = info[1], info[2]
                    cur = var.get()
                    # normaliza a inglés canonical y re-traduce a display target
                    can = _opt_to_en(cur)
                    disp = _opt_to_es(can) if self.lang == "es" else can
                    if disp != cur:
                        var.set(disp)
        except Exception:
            pass

    # === End I18N ===

    @staticmethod
    def _put(mapping: dict, names, value):
        for n in names:
            mapping[f"{{{{{n}}}}}"] = str(value)

    def _pack_generate(self):
        # Uso del backend reutilizable para evitar duplicar lógica de negocio.
        plant = self.var_plant.get(); ruta = os.path.join(_app_dir(), plant)
        if not os.path.exists(ruta):
            messagebox.showerror(APP_TITLE, f"No se encontró la plantilla: {plant}"); return

        overrides = {}
        for name, info in self.opt_vars.items():
            if info[0] == "chk":
                var = info[1]
                overrides[name] = bool(var.get())
            else:
                var = info[1]
                overrides[name] = var.get()

        for i in range(1, 4):
            overrides[f"contrato{i}_porcentaje"] = self.var_conc_pct[i-1].get()
            overrides[f"contrato{i}_condicion"] = self.var_conc_venc[i-1].get()

        out_docx = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word","*.docx")],
            initialfile=f"Cotizacion_{self.var_cliente.get()}_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        )
        if not out_docx:
            return
        out_pdf = out_docx.replace(".docx", ".pdf")

        try:
            # Delegar la generación al backend reutilizable
            validez_txt = self.var_validez.get()
            try:
                validez_num = int("".join(ch for ch in validez_txt if ch.isdigit())) if validez_txt else None
            except Exception:
                validez_num = None
            resultado = generar_cotizacion_backend(
                modelo=plant,
                cliente=self.var_cliente.get(),
                asesor=self.var_asesor.get(),
                fecha=self.var_fecha.get(),
                validez_dias=validez_num,
                moneda=None,
                flete_texto=None,
                flete_monto=None,
                notas=None,
                opciones_overrides=overrides,
                ruta_salida_word=out_docx,
                ruta_salida_pdf=out_pdf,
            )
            self.var_total.set(_fmt_money(Decimal(resultado.get("total", 0)), resultado.get("moneda", "USD")))
            stored_docx, stored_pdf = self._history_store_files("packaging", resultado.get("ruta_word", ""), resultado.get("ruta_pdf", ""))
            history_entry = {
                "fecha": datetime.now().strftime("%Y-%m-%d %H:%M"),
                "cliente": self.var_cliente.get(),
                "plantilla": plant,
                "monto": _fmt_money(Decimal(resultado.get("total", 0)), resultado.get("moneda", "USD")),
                "docx": stored_docx or resultado.get("ruta_word", ""),
                "pdf": stored_pdf or resultado.get("ruta_pdf", ""),
                "total_numeric": resultado.get("total", 0),
                "original_docx": resultado.get("ruta_word", ""),
                "original_pdf": resultado.get("ruta_pdf", ""),
            }
            self._history_add("packaging", history_entry)
            messagebox.showinfo(APP_TITLE, "Documento generado.")
        except Exception as e:
            _write_error_log(e)
            messagebox.showerror(APP_TITLE, f"Error al generar:\n{e}\n\nRevise error.log para más detalles.")

    @staticmethod
    def _convert_docx_to_pdf(input_path: str, output_path: str) -> str:
        try:
            out_dir = os.path.dirname(output_path) or "."
            os.makedirs(out_dir, exist_ok=True)
            subprocess.run([
                "libreoffice", "--headless", "--convert-to", "pdf",
                "--outdir", out_dir, input_path
            ], check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
            gen = os.path.join(out_dir, os.path.splitext(os.path.basename(input_path))[0] + ".pdf")
            if os.path.abspath(gen) != os.path.abspath(output_path):
                try: os.replace(gen, output_path)
                except Exception:
                    shutil.copyfile(gen, output_path); os.remove(gen)
            return output_path if os.path.exists(output_path) else ""
        except Exception:
            return ""

    # ===== Materials =====
    def _build_materials(self, parent):
        head = ttk.LabelFrame(parent, text="Encabezado"); head.pack(fill="x", padx=8, pady=6)
        self._register_translated_widget("Encabezado", head)
        # Materials-specific logo
        try:
            import os
            from tkinter import PhotoImage
            mat_logo_path = os.path.join(_app_dir(), "logo_materials.png")
            if os.path.exists(mat_logo_path):
                img = tk.PhotoImage(file=mat_logo_path)
                # downscale to fit max height 80px
                try:
                    h = img.height(); w = img.width()
                    factor = max(1, int((h/80)+0.999))
                    if factor > 1:
                        img = img.subsample(factor, factor)
                except Exception:
                    pass
                self._logo_materials = img
            else:
                self._logo_materials = self._logo_img
            if self._logo_materials:
                ttk.Label(head, image=self._logo_materials).grid(row=0, column=5, rowspan=2, padx=(12,0), pady=2, sticky="e")
        except Exception:
            pass

        self.var_m_cliente = tk.StringVar(); self.var_m_fecha = tk.StringVar(value=datetime.now().strftime("%d/%m/%Y"))
        self.var_m_asesor = tk.StringVar(value="José Manuel Rayotorres Martínez"); self.var_m_validez = tk.StringVar(value="30 días")
        lbl_mc = ttk.Label(head, text="Cliente:"); lbl_mc.grid(row=0, column=0, padx=4, pady=2, sticky="e"); self._register_translated_widget("Cliente:", lbl_mc)
        ttk.Entry(head, textvariable=self.var_m_cliente, width=38).grid(row=0, column=1, padx=4, pady=2, sticky="w")
        lbl_mf = ttk.Label(head, text="Fecha:"); lbl_mf.grid(row=0, column=2, padx=4, pady=2, sticky="e"); self._register_translated_widget("Fecha:", lbl_mf)
        ttk.Entry(head, textvariable=self.var_m_fecha, width=12).grid(row=0, column=3, padx=4, pady=2, sticky="w")
        lbl_ma = ttk.Label(head, text="Asesor:"); lbl_ma.grid(row=1, column=0, padx=4, pady=2, sticky="e"); self._register_translated_widget("Asesor:", lbl_ma)
        ttk.Entry(head, textvariable=self.var_m_asesor, width=38).grid(row=1, column=1, padx=4, pady=2, sticky="w")
        lbl_mv = ttk.Label(head, text="Validez:"); lbl_mv.grid(row=1, column=2, padx=4, pady=2, sticky="e"); self._register_translated_widget("Validez:", lbl_mv)
        ttk.Entry(head, textvariable=self.var_m_validez, width=12).grid(row=1, column=3, padx=4, pady=2, sticky="w")
        # Selector de plantilla
        from glob import glob
        default_tpl = os.path.join(_app_dir(), "Cotizacion Materials.docx")
        choices = []
        try:
            choices = [p for p in glob(os.path.join(_app_dir(), "*.docx")) if "Cotizacion" in os.path.basename(p)]
        except Exception:
            choices = []
        if default_tpl not in choices:
            choices = [default_tpl] + choices
        self.var_m_tpl = tk.StringVar(value=choices[0] if choices else default_tpl)
        lbl_tpl = ttk.Label(head, text="Plantilla:"); lbl_tpl.grid(row=2, column=0, padx=4, pady=2, sticky="e"); self._register_translated_widget("Plantilla:", lbl_tpl)
        tpl_box = ttk.Frame(head); tpl_box.grid(row=2, column=1, columnspan=3, padx=4, pady=2, sticky="we")
        tpl_box.columnconfigure(0, weight=1)
        self.cbo_m_tpl = ttk.Combobox(tpl_box, textvariable=self.var_m_tpl, values=choices, state="readonly")
        self.cbo_m_tpl.grid(row=0, column=0, sticky="we")
        def _browse_tpl():
            path = filedialog.askopenfilename(title="Seleccionar plantilla DOCX", filetypes=[("Word","*.docx")], initialdir=_app_dir())
            if path:
                self.var_m_tpl.set(path)
                vals = list(self.cbo_m_tpl["values"])
                if path not in vals:
                    self.cbo_m_tpl["values"] = vals + [path]
        btn_tpl = ttk.Button(tpl_box, text="Buscar plantilla", command=_browse_tpl)
        btn_tpl.grid(row=0, column=1, padx=(6,0))
        self._register_translated_widget("Buscar plantilla", btn_tpl)

        # Dynamic items area
        items_box = ttk.LabelFrame(parent, text="Ítems"); self._register_translated_widget("Ítems", items_box)
        items_box.pack(fill="x", padx=8, pady=6)
        self.items_container = ttk.Frame(items_box); self.items_container.pack(fill="x", padx=6, pady=4)
        self.item_rows = []

        hdr = ttk.Frame(items_box); hdr.pack(fill="x", padx=6)
        lbl_prod = ttk.Label(hdr, text="Producto", width=60, anchor="w")
        lbl_prod.pack(side="left", padx=(0,6))
        self._register_translated_widget("Producto", lbl_prod)
        lbl_qty = ttk.Label(hdr, text="Cantidad", width=6, anchor="w")
        lbl_qty.pack(side="left", padx=(0,6))
        self._register_translated_widget("Cantidad", lbl_qty)
        lbl_unit = ttk.Label(hdr, text="Precio unitario", width=14, anchor="w")
        lbl_unit.pack(side="left", padx=(0,6))
        self._register_translated_widget("Precio unitario", lbl_unit)
        def add_row(desc="", qty="1", each="0"):
            row = ttk.Frame(self.items_container); row.pack(fill="x", pady=2)
            v_desc, v_qty, v_each = tk.StringVar(value=desc), tk.StringVar(value=qty), tk.StringVar(value=each)
            e1 = ttk.Entry(row, textvariable=v_desc, width=60); e1.pack(side="left", padx=(0,6))
            e2 = ttk.Entry(row, textvariable=v_qty, width=6); e2.pack(side="left", padx=(0,6))
            e3 = ttk.Entry(row, textvariable=v_each, width=10); e3.pack(side="left", padx=(0,6))
            bdel = ttk.Button(row, text="–", width=3, command=lambda: remove_row(row))
            bdel.pack(side="left")
            for v in (v_desc, v_qty, v_each):
                v.trace_add("write", lambda *_: self._mat_update_totals())
            self.item_rows.append((row, v_desc, v_qty, v_each))
            self._mat_update_totals()

        def remove_row(roww):
            self.item_rows = [t for t in self.item_rows if t[0] is not roww]
            try: roww.destroy()
            except Exception: pass
            self._mat_update_totals()

        btns = ttk.Frame(items_box); btns.pack(fill="x", padx=6, pady=2)
        btn_addi = ttk.Button(btns, text="Agregar ítem", command=lambda: add_row()); btn_addi.pack(side="left"); self._register_translated_widget("Agregar ítem", btn_addi)
        btn_delli = ttk.Button(btns, text="Eliminar último", command=lambda: (self.item_rows and remove_row(self.item_rows[-1][0]))); btn_delli.pack(side="left", padx=6); self._register_translated_widget("Eliminar último", btn_delli)

        # Opciones
        opts = ttk.LabelFrame(parent, text="Opciones"); opts.pack(fill="x", padx=8, pady=6)
        self._register_translated_widget("Opciones", opts)
        self.var_moneda = tk.StringVar(value="MXN")
        self.var_iva = tk.StringVar(value="16")
        lbl_mon = ttk.Label(opts, text="Moneda"); lbl_mon.grid(row=0, column=0, padx=4, pady=2); self._register_translated_widget("Moneda", lbl_mon); ttk.Combobox(opts, textvariable=self.var_moneda, values=["USD","MXN"], width=6, state="readonly").grid(row=0, column=1, padx=4, pady=2)
        lbl_iva_pct = ttk.Label(opts, text="IVA (%)"); lbl_iva_pct.grid(row=0, column=2, padx=4, pady=2); self._register_translated_widget("IVA (%)", lbl_iva_pct)
        cb_iva = ttk.Combobox(opts, textvariable=self.var_iva, values=[str(x) for x in range(0,33,1)], width=5, state="readonly")
        cb_iva.grid(row=0, column=3, padx=4, pady=2); cb_iva.bind("<<ComboboxSelected>>", lambda e: self._mat_update_totals())

        fle = ttk.LabelFrame(parent, text="Flete y notas"); fle.pack(fill="x", padx=8, pady=4)
        self._register_translated_widget("Flete y notas", fle)
        self.var_m_flete_modo = tk.StringVar(value="No incluido"); self.var_m_flete_monto = tk.StringVar(value="0")
        lbl_fle = ttk.Label(fle, text="Flete"); lbl_fle.grid(row=0, column=0, padx=4, pady=2, sticky="w"); self._register_translated_widget("Flete", lbl_fle)
        flete_values = ["No aplica","Incluido","No incluido"]
        cb_flete = ttk.Combobox(fle, textvariable=self.var_m_flete_modo, values=flete_values, width=12, state="readonly")
        cb_flete.grid(row=0, column=1, padx=4, pady=2, sticky="w"); cb_flete.bind("<<ComboboxSelected>>", lambda e: self._mat_update_totals())
        try:
            cb_flete.current(flete_values.index("No incluido"))
        except ValueError:
            pass
        self._register_translated_combobox(cb_flete, self.var_m_flete_modo, flete_values)
        lbl_monto = ttk.Label(fle, text="Monto"); lbl_monto.grid(row=0, column=2, padx=4, pady=2, sticky="w"); self._register_translated_widget("Monto", lbl_monto)
        ent_flete = ttk.Entry(fle, textvariable=self.var_m_flete_monto, width=12)
        ent_flete.grid(row=0, column=3, padx=4, pady=2, sticky="w"); ent_flete.bind("<FocusOut>", lambda e: (self._normalize_money(self.var_m_flete_monto), self._mat_update_totals()))
        self.txt_m_notas = tk.Text(fle, width=80, height=2)
        lbl_notas = ttk.Label(fle, text="Notas"); lbl_notas.grid(row=1, column=0, padx=4, pady=2, sticky="nw"); self._register_translated_widget("Notas", lbl_notas); self.txt_m_notas.grid(row=1, column=1, columnspan=3, padx=4, pady=2, sticky="we")

        bar = ttk.Frame(parent); bar.pack(fill="x", padx=8, pady=(6,2))
        self.lbl_sub = ttk.Label(bar, text="Subtotal: US$0.00", font=("Segoe UI", 10)); self.lbl_sub.pack(side="left", padx=(0,16))
        self.lbl_iva = ttk.Label(bar, text="IVA: US$0.00", font=("Segoe UI", 10)); self.lbl_iva.pack(side="left", padx=(0,16))
        self.lbl_tot = ttk.Label(bar, text="Total: US$0.00", font=("Segoe UI", 10, "bold")); self.lbl_tot.pack(side="left")
        btn_gen = ttk.Button(bar, text="Generar Word/PDF", command=self._mat_generate)
        btn_gen.pack(side="right")
        self._register_translated_widget("Generar Word/PDF", btn_gen)

        # start with one row
        add_row()

    def _mat_update_totals(self):
        subtotal = Decimal("0")
        # sum from dynamic rows
        for (_row, v_desc, v_qty, v_each) in getattr(self, "item_rows", []):
            try:
                q = _parse_decimal_safe(v_qty.get())
                e = _parse_decimal_safe(v_each.get())
                subtotal += q * e
            except Exception:
                continue
        iva_pct = _parse_decimal_safe(self.var_iva.get())
        iva = (subtotal * iva_pct) / Decimal("100")
        total = subtotal + iva
        self._mat_totals = {"subtotal": subtotal, "iva": iva, "total": total, "iva_pct": iva_pct}
        self._update_material_totals_labels()

def main():
    cfg = _read_cfg()
    try:
        app = App(cfg)
        app.mainloop()
    except Exception as exc:  # pragma: no cover - defensive
        _write_error_log(exc)
        try:
            root = tk.Tk()
            root.withdraw()
            messagebox.showerror(
                APP_TITLE,
                "Ocurrió un error y la aplicación se cerró.\n"
                "Revisa 'error.log' en la carpeta del programa para más detalles.\n\n"
                f"Detalle: {exc}"
            )
        except Exception:
            pass

if __name__ == "__main__":
    main()
