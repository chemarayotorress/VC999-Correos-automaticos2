"""DOCX template placeholder mapping utilities."""
from __future__ import annotations

import json
import os
import re
import zipfile
from dataclasses import dataclass
from typing import Dict, List, Optional, Set, Tuple, Union

try:  # Optional dependency - la app debe iniciar aunque falte python-docx
    from docx import Document  # type: ignore
except Exception:  # pragma: no cover - entorno sin python-docx
    Document = None  # type: ignore

try:
    import tkinter as tk
    from tkinter import ttk, messagebox, filedialog
except Exception:  # pragma: no cover
    tk = None  # type: ignore
    ttk = None  # type: ignore
    messagebox = None  # type: ignore
    filedialog = None  # type: ignore

TEMPLATE_MAP_FILE = "template_mappings.json"
APP_TITLE = "VC999 Cotizador"
_WARNED_DOCX_MISSING = False


@dataclass
class MappingOption:
    key: str
    label: str
    category: str = ""


class TemplateMappingManager:
    @staticmethod
    def _load_raw() -> Dict[str, Dict[str, Dict[str, Dict[str, str]]]]:
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), TEMPLATE_MAP_FILE)
        if not os.path.exists(path):
            return {}
        try:
            with open(path, "r", encoding="utf-8") as f:
                data = json.load(f)
                if isinstance(data, dict):
                    return data
        except Exception:
            pass
        return {}

    @staticmethod
    def _save_raw(data: Dict[str, Dict[str, Dict[str, Dict[str, str]]]]):
        path = os.path.join(os.path.dirname(os.path.abspath(__file__)), TEMPLATE_MAP_FILE)
        with open(path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)

    @classmethod
    def load_mapping(cls, kind: str, template_name: str) -> Dict[str, Dict[str, str]]:
        data = cls._load_raw()
        return data.get(kind, {}).get(template_name, {})

    @classmethod
    def list_templates(cls, kind: Optional[str] = None) -> Union[Dict[str, List[str]], List[str]]:
        data = cls._load_raw()
        if kind is None:
            return {k: sorted(v.keys()) for k, v in data.items()}
        return sorted(data.get(kind, {}).keys())

    @classmethod
    def save_mapping(cls, kind: str, template_name: str, mapping: Dict[str, Dict[str, str]]):
        data = cls._load_raw()
        data.setdefault(kind, {})[template_name] = mapping
        cls._save_raw(data)

    @staticmethod
    def extract_placeholders(docx_path: str) -> List[str]:
        placeholders: List[str] = []
        try:
            with zipfile.ZipFile(docx_path, "r") as zf:
                for name in zf.namelist():
                    if not name.lower().endswith(".xml"):
                        continue
                    try:
                        xml = zf.read(name).decode("utf-8")
                    except Exception:
                        continue
                    placeholders.extend(re.findall(r"\{\{[^{}]+\}\}", xml))
        except Exception:
            return []
        cleaned: List[str] = []
        seen = set()
        for ph in placeholders:
            ph = ph.strip()
            if ph not in seen:
                cleaned.append(ph)
                seen.add(ph)
        return cleaned


def _normalize_placeholder(placeholder: str) -> str:
    if not isinstance(placeholder, str):
        return ""
    text = placeholder.strip()
    if text.startswith("{{") and text.endswith("}}"):
        text = text[2:-2]
    return text.strip()


def ensure_placeholders(docx_path: str, options: List[MappingOption]) -> Set[str]:
    """Ensure the DOCX contains placeholders for every mapping option.

    Returns a set with the placeholders (including braces) that were added.
    """
    added: Set[str] = set()
    global _WARNED_DOCX_MISSING
    if Document is None:
        if not _WARNED_DOCX_MISSING and messagebox:
            _WARNED_DOCX_MISSING = True
            try:
                messagebox.showwarning(
                    APP_TITLE,
                    "No se pueden agregar campos automáticamente porque falta la"
                    " dependencia 'python-docx'. Instálala y vuelve a intentar.",
                )
            except Exception:
                pass
        return added
    if not docx_path or not os.path.exists(docx_path):
        return added
    try:
        existing = {
            _normalize_placeholder(ph)
            for ph in TemplateMappingManager.extract_placeholders(docx_path)
        }
        doc = Document(docx_path)
    except Exception as exc:
        if messagebox:
            try:
                messagebox.showerror(
                    APP_TITLE,
                    "No se pudo preparar la plantilla para agregar marcadores."
                    f"\n\nDetalle: {exc}",
                )
            except Exception:
                pass
        return added

    required: List[Tuple[str, str]] = []
    for option in options:
        key = (option.key or "").strip()
        if not key:
            continue
        placeholder = f"{{{{{key}}}}}"
        if _normalize_placeholder(placeholder) not in existing:
            required.append((placeholder, option.label or key))
            added.add(placeholder)

    if not required:
        return added

    note_text = (
        "Marcadores automáticos agregados por VC999 Cotizador (puedes moverlos donde correspondan)."
    )
    has_note = any(p.text.strip() == note_text for p in doc.paragraphs)
    if not has_note:
        if doc.paragraphs and doc.paragraphs[-1].text.strip():
            doc.add_paragraph("")
        doc.add_paragraph(note_text)

    target_table = None
    for table in doc.tables:
        try:
            header_0 = table.rows[0].cells[0].text.strip().lower()
            header_1 = table.rows[0].cells[1].text.strip().lower()
        except Exception:
            continue
        if header_0 == "campo" and "marcador" in header_1:
            target_table = table
            break

    if target_table is None:
        target_table = doc.add_table(rows=1, cols=2)
        target_table.autofit = True
        target_table.rows[0].cells[0].text = "Campo"
        target_table.rows[0].cells[1].text = "Marcador"

    existing_markers = {
        cell.text.strip()
        for row in target_table.rows[1:]
        for cell in row.cells[1:2]
        if cell.text
    }

    for placeholder, label in required:
        if placeholder in existing_markers:
            continue
        row = target_table.add_row()
        row.cells[0].text = label
        row.cells[1].text = placeholder

    try:
        doc.save(docx_path)
    except Exception as exc:
        if messagebox:
            try:
                messagebox.showerror(
                    APP_TITLE,
                    "No se pudieron guardar los marcadores automáticamente."
                    " Verifica que el archivo no esté abierto en Word e inténtalo"
                    f" de nuevo.\n\nDetalle: {exc}",
                )
            except Exception:
                pass
        return set()

    return added


